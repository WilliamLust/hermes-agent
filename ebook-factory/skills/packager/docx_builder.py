#!/usr/bin/env python3
"""
DOCX Builder — python-docx native generation for Ebook Factory
================================================================

Builds KDP-ready DOCX files directly from markdown chapter sources,
replacing the pandoc HTML→DOCX pipeline. This eliminates pandoc
artifacts (duplicate paragraphs, broken styles) and gives full
control over every style, paragraph, and run.

Usage (standalone):
    python3 docx_builder.py --book-dir ~/.hermes/ebook-factory/workbooks/book-my-topic/

Usage (imported):
    from docx_builder import build_docx
    docx_path = build_docx(meta, chapters, output_dir)

Input:
    - chapters: list of dicts with keys {number, title, content (markdown string)}
    - meta: dict with keys {title, subtitle, author, slug, date, language}

Output:
    - Path to the generated .docx file
"""

import re
import sys
import os
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


# ============================================================================
# LOGGING
# ============================================================================

def log(msg: str) -> None:
    ts = datetime.now().strftime("%H:%M:%S")
    print(f"[{ts}] {msg}", flush=True)


# ============================================================================
# STYLE SETUP
# ============================================================================

GARAMOND = "Garamond"


def _setup_styles(doc: Document) -> None:
    """
    Create or update all required Word styles in the document.

    Styles:
        Normal        — body text, Garamond 11pt, left-aligned, 0.25" first-line indent
        Heading 1     — chapter titles, Garamond Bold 18pt, centered, page break before
        Heading 2     — section headings, Garamond Bold 14pt, left-aligned, keep with next
        Heading 3     — subsections, Garamond Bold Italic 12pt, left-aligned, keep with next
        Heading 4     — sub-subsections, Garamond Bold 11pt, left-aligned, keep with next
        Quote         — blockquotes, Garamond Italic 10.5pt, 0.5" left/right indent
        List Paragraph — list items, Garamond 11pt, hanging indent
        Copyright     — 8pt, left-aligned, single spacing
    """
    styles = doc.styles

    # --- Normal (body text) ---
    normal = styles['Normal']
    normal.font.name = GARAMOND
    normal.font.size = Pt(11)
    normal.font.bold = False
    normal.font.italic = False
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    normal.paragraph_format.first_line_indent = Inches(0.25)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    # Set the theme font for Garamond (rFonts for ascii and hAnsi)
    _set_theme_fonts(normal.font, GARAMOND)

    # --- Heading 1 (chapter titles) ---
    h1 = styles['Heading 1']
    h1.font.name = GARAMOND
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.italic = False
    h1.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(24)
    h1.paragraph_format.page_break_before = True
    h1.paragraph_format.keep_with_next = False
    h1.paragraph_format.first_line_indent = None  # no indent on headings
    _set_theme_fonts(h1.font, GARAMOND)
    _set_outline_level(h1, 1)

    # --- Heading 2 (section headings) ---
    h2 = styles['Heading 2']
    h2.font.name = GARAMOND
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.italic = False
    h2.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.space_before = Pt(24)
    h2.paragraph_format.space_after = Pt(12)
    h2.paragraph_format.keep_with_next = True
    h2.paragraph_format.first_line_indent = None
    _set_theme_fonts(h2.font, GARAMOND)
    _set_outline_level(h2, 2)

    # --- Heading 3 (subsections) ---
    h3 = styles['Heading 3']
    h3.font.name = GARAMOND
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.italic = True
    h3.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3.paragraph_format.space_before = Pt(18)
    h3.paragraph_format.space_after = Pt(6)
    h3.paragraph_format.keep_with_next = True
    h3.paragraph_format.first_line_indent = None
    _set_theme_fonts(h3.font, GARAMOND)
    _set_outline_level(h3, 3)

    # --- Heading 4 (sub-subsections) ---
    h4 = styles['Heading 4']
    h4.font.name = GARAMOND
    h4.font.size = Pt(11)
    h4.font.bold = True
    h4.font.italic = False
    h4.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    h4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h4.paragraph_format.space_before = Pt(12)
    h4.paragraph_format.space_after = Pt(6)
    h4.paragraph_format.keep_with_next = True
    h4.paragraph_format.first_line_indent = None
    _set_theme_fonts(h4.font, GARAMOND)
    _set_outline_level(h4, 4)

    # --- Quote (blockquotes) ---
    # Use built-in 'Quote' style if it exists, otherwise create it
    quote_style = _get_or_create_style(styles, 'Quote', WD_STYLE_TYPE.PARAGRAPH)
    quote_style.font.name = GARAMOND
    quote_style.font.size = Pt(10.5)
    quote_style.font.bold = False
    quote_style.font.italic = True
    quote_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    quote_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    quote_style.paragraph_format.left_indent = Inches(0.5)
    quote_style.paragraph_format.right_indent = Inches(0.5)
    quote_style.paragraph_format.space_before = Pt(12)
    quote_style.paragraph_format.space_after = Pt(12)
    quote_style.paragraph_format.first_line_indent = None
    _set_theme_fonts(quote_style.font, GARAMOND)

    # --- List Paragraph ---
    list_para = _get_or_create_style(styles, 'List Paragraph', WD_STYLE_TYPE.PARAGRAPH)
    list_para.font.name = GARAMOND
    list_para.font.size = Pt(11)
    list_para.font.bold = False
    list_para.font.italic = False
    list_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    list_para.paragraph_format.space_before = Pt(0)
    list_para.paragraph_format.space_after = Pt(0)
    list_para.paragraph_format.first_line_indent = None
    _set_theme_fonts(list_para.font, GARAMOND)

    # --- Copyright ---
    copyright_style = _get_or_create_style(styles, 'Copyright', WD_STYLE_TYPE.PARAGRAPH)
    copyright_style.font.name = GARAMOND
    copyright_style.font.size = Pt(8)
    copyright_style.font.bold = False
    copyright_style.font.italic = False
    copyright_style.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    copyright_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    copyright_style.paragraph_format.space_before = Pt(0)
    copyright_style.paragraph_format.space_after = Pt(0)
    copyright_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    copyright_style.paragraph_format.first_line_indent = None
    _set_theme_fonts(copyright_style.font, GARAMOND)


def _set_theme_fonts(font, name: str) -> None:
    """Set the rFonts element to ensure the font name sticks across Word implementations."""
    try:
        rpr = font.element
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is None:
            rfonts = OxmlElement('w:rFonts')
            rpr.insert(0, rfonts)
        rfonts.set(qn('w:ascii'), name)
        rfonts.set(qn('w:hAnsi'), name)
        rfonts.set(qn('w:cs'), name)
    except Exception:
        pass  # Non-critical; the font.name property is the primary setter


def _set_outline_level(style, level: int) -> None:
    """Set the outline level on a heading style via the XML pPr element."""
    try:
        ppr = style.element.find(qn('w:pPr'))
        if ppr is None:
            ppr = OxmlElement('w:pPr')
            style.element.insert(0, ppr)
        outline = ppr.find(qn('w:outlineLvl'))
        if outline is None:
            outline = OxmlElement('w:outlineLvl')
            ppr.append(outline)
        outline.set(qn('w:val'), str(level))
    except Exception:
        pass


def _get_or_create_style(styles, name: str, style_type):
    """Get an existing style or create a new one."""
    try:
        style = styles[name]
        return style
    except KeyError:
        return styles.add_style(name, style_type)


# ============================================================================
# BOOKMARK HELPERS
# ============================================================================

_next_bookmark_id = 100  # start above any built-in IDs


def _add_bookmark(paragraph, name: str) -> None:
    """
    Add a Word bookmark (w:bookmarkStart / w:bookmarkEnd) to a paragraph.
    This allows TOC hyperlinks and KDP detection of the TOC.
    """
    global _next_bookmark_id
    bm_id = _next_bookmark_id
    _next_bookmark_id += 1

    # bookmarkStart goes at the beginning of the paragraph
    bm_start = OxmlElement('w:bookmarkStart')
    bm_start.set(qn('w:id'), str(bm_id))
    bm_start.set(qn('w:name'), name)
    paragraph._element.insert(0, bm_start)

    # bookmarkEnd goes at the end
    bm_end = OxmlElement('w:bookmarkEnd')
    bm_end.set(qn('w:id'), str(bm_id))
    paragraph._element.append(bm_end)


def _add_hyperlink(paragraph, bookmark_name: str, text: str) -> None:
    """
    Add an internal hyperlink (to a bookmark) within a paragraph.
    Uses w:hyperlink with w:anchor attribute.
    """
    # Create the hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), bookmark_name)

    # Create the run inside the hyperlink
    run_elem = OxmlElement('w:r')

    # Run properties (font, size, color)
    rpr = OxmlElement('w:rPr')
    rfonts = OxmlElement('w:rFonts')
    rfonts.set(qn('w:ascii'), GARAMOND)
    rfonts.set(qn('w:hAnsi'), GARAMOND)
    rfonts.set(qn('w:cs'), GARAMOND)
    rpr.append(rfonts)

    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(22))  # 11pt = 22 half-points
    rpr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(22))
    rpr.append(szCs)

    color = OxmlElement('w:color')
    color.set(qn('w:val'), '444444')
    rpr.append(color)

    # Underline (hyperlinks typically underlined)
    # We skip underline for cleaner TOC look; KDP doesn't require it
    # u = OxmlElement('w:u')
    # u.set(qn('w:val'), 'single')
    # rpr.append(u)

    run_elem.append(rpr)

    # Text
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    run_elem.append(t)

    hyperlink.append(run_elem)
    paragraph._element.append(hyperlink)


# ============================================================================
# DOT LEADER TAB STOP (for TOC)
# ============================================================================

def _add_dot_leader_tab(paragraph, tab_position_inches: float = 6.0) -> None:
    """
    Add a right-aligned tab stop with dot leaders to a paragraph.
    Used in TOC entries to create "Chapter Title ..... page" layout.
    """
    ppr = paragraph._element.find(qn('w:pPr'))
    if ppr is None:
        ppr = OxmlElement('w:pPr')
        paragraph._element.insert(0, ppr)

    tabs = ppr.find(qn('w:tabs'))
    if tabs is None:
        tabs = OxmlElement('w:tabs')
        ppr.append(tabs)

    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:leader'), 'dot')
    tab.set(qn('w:pos'), str(Inches(tab_position_inches)))
    tabs.append(tab)


# ============================================================================
# MARKDOWN PARSER — inline formatting
# ============================================================================

# Regex patterns for inline formatting, ordered by specificity
# ***bold italic*** must be matched before **bold** and *italic*
_INLINE_RE = re.compile(
    r'(\*\*\*(.+?)\*\*\*)'    # ***bold italic***
    r'|(\*\*(.+?)\*\*)'       # **bold**
    r'|(\*(.+?)\*)'           # *italic*
    r'|(`(.+?)`)'             # `inline code`
    r'|(\[(.+?)\]\((.+?)\))' # [link text](url)
)


def _parse_inline(text: str) -> list:
    """
    Parse inline markdown formatting and return a list of
    (text, bold, italic, code, link_url) tuples for each segment.

    Handles:
        ***bold italic***  -> bold=True, italic=True
        **bold**           -> bold=True
        *italic*           -> italic=True
        `inline code`      -> code=True
        [text](url)        -> link_url=url (text only in DOCX)
    """
    segments = []
    last_end = 0

    for m in _INLINE_RE.finditer(text):
        # Add any plain text before this match
        if m.start() > last_end:
            plain = text[last_end:m.start()]
            if plain:
                segments.append((plain, False, False, False, None))

        if m.group(1):  # ***bold italic***
            segments.append((m.group(2), True, True, False, None))
        elif m.group(3):  # **bold**
            segments.append((m.group(4), True, False, False, None))
        elif m.group(5):  # *italic*
            segments.append((m.group(6), False, True, False, None))
        elif m.group(7):  # `code`
            segments.append((m.group(8), False, False, True, None))
        elif m.group(9):  # [text](url)
            link_text = m.group(10)
            link_url = m.group(11)
            # For DOCX we render the link as text (url in parens for readability)
            segments.append((f"{link_text} ({link_url})", False, False, False, None))

        last_end = m.end()

    # Remaining text
    if last_end < len(text):
        remaining = text[last_end:]
        if remaining:
            segments.append((remaining, False, False, False, None))

    # If no inline formatting found, return the whole text as plain
    if not segments:
        segments.append((text, False, False, False, None))

    return segments


def _add_inline_runs(paragraph, text: str, base_font_name: str = GARAMOND,
                     base_font_size: Pt = None, base_color: RGBColor = None) -> None:
    """
    Parse inline markdown formatting in text and add runs to the paragraph.
    Preserves bold, italic, bold-italic, and inline code at the run level.
    """
    segments = _parse_inline(text)

    for seg_text, bold, italic, code, link_url in segments:
        if not seg_text:
            continue

        run = paragraph.add_run(seg_text)

        # Apply base font
        run.font.name = base_font_name
        _set_run_fonts(run, base_font_name)

        if base_font_size is not None:
            run.font.size = base_font_size

        if base_color is not None:
            run.font.color.rgb = base_color

        # Apply inline formatting
        if code:
            run.font.name = "Consolas"
            _set_run_fonts(run, "Consolas")
            run.font.size = Pt(9.5) if base_font_size is None else Pt(base_font_size.pt - 1.5)
        else:
            run.font.bold = bold
            run.font.italic = italic


def _set_run_fonts(run, name: str) -> None:
    """Set rFonts on a run element."""
    try:
        rpr = run._element.find(qn('w:rPr'))
        if rpr is None:
            rpr = OxmlElement('w:rPr')
            run._element.insert(0, rpr)
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is None:
            rfonts = OxmlElement('w:rFonts')
            rpr.insert(0, rfonts)
        rfonts.set(qn('w:ascii'), name)
        rfonts.set(qn('w:hAnsi'), name)
        rfonts.set(qn('w:cs'), name)
    except Exception:
        pass


# ============================================================================
# MARKDOWN TO DOCX BLOCK PARSER
# ============================================================================

def parse_markdown_to_docx(doc: Document, md_text: str,
                           first_para_after_heading: bool = False) -> bool:
    """
    Parse a markdown string and add properly formatted paragraphs/runs
    to the document.

    Handles:
        - H1-H4 headings
        - **bold**, *italic*, ***bold italic*** inline
        - `inline code`
        - [links](url) rendered as text
        - Numbered lists (1. 2. 3.)
        - Bullet lists (- or * or +)
        - Blockquotes (> prefix)
        - Paragraph breaks (blank lines)
        - First-paragraph-after-heading: no first-line indent (flush left)

    Args:
        doc: The python-docx Document object
        md_text: Markdown content string
        first_para_after_heading: If True, the first paragraph will have
            no first-line indent (flush left). Typically True when called
            right after adding a heading.

    Returns:
        True if the last element added was a heading (caller should set
        first_para_after_heading for the next call), False otherwise.
    """
    # Strip validation report and HTML comments
    md_text = re.sub(r'\n## Validation Report.*', '', md_text, flags=re.DOTALL)
    md_text = re.sub(r'<!--.*?-->', '', md_text, flags=re.DOTALL)

    lines = md_text.split('\n')

    # State tracking
    in_blockquote = False
    blockquote_lines = []
    in_bullet_list = False
    in_numbered_list = False
    pending_para_lines = []
    just_added_heading = first_para_after_heading
    first_body_para_done = not first_para_after_heading  # already past first-para if False

    def flush_paragraph():
        """Flush accumulated paragraph lines into a Normal paragraph."""
        nonlocal pending_para_lines, just_added_heading, first_body_para_done
        if not pending_para_lines:
            return
        text = ' '.join(pending_para_lines).strip()
        pending_para_lines = []
        if not text:
            return

        p = doc.add_paragraph(style='Normal')

        # First paragraph after heading: no first-line indent
        if just_added_heading and not first_body_para_done:
            p.paragraph_format.first_line_indent = Pt(0)
            first_body_para_done = True
        else:
            p.paragraph_format.first_line_indent = Inches(0.25)

        _add_inline_runs(p, text)
        just_added_heading = False

    def flush_blockquote():
        """Flush accumulated blockquote lines into a Quote paragraph."""
        nonlocal blockquote_lines, just_added_heading, first_body_para_done
        if not blockquote_lines:
            return
        text = ' '.join(blockquote_lines).strip()
        blockquote_lines = []
        if not text:
            return

        p = doc.add_paragraph(style='Quote')
        _add_inline_runs(p, text)
        just_added_heading = False

    def close_lists():
        """Close any open list context."""
        nonlocal in_bullet_list, in_numbered_list
        in_bullet_list = False
        in_numbered_list = False

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # --- Blockquote handling ---
        # Detect blockquote start: line begins with >
        if stripped.startswith('>'):
            # If we were in a paragraph, flush it
            flush_paragraph()
            close_lists()

            # Accumulate blockquote content (strip > prefix)
            bq_text = re.sub(r'^>\s?', '', stripped)
            blockquote_lines.append(bq_text)
            in_blockquote = True
            i += 1
            continue

        # If we were in a blockquote and this line is not a blockquote, flush it
        if in_blockquote and not stripped.startswith('>'):
            flush_blockquote()
            in_blockquote = False

        # --- Empty line = paragraph break ---
        if not stripped:
            flush_paragraph()
            # Close lists on blank line
            close_lists()
            i += 1
            continue

        # --- Headings (H1-H4) ---
        heading_match = re.match(r'^(#{1,4})\s+(.+)$', stripped)
        if heading_match:
            flush_paragraph()
            close_lists()

            level = len(heading_match.group(1))
            heading_text = heading_match.group(2).strip()

            style_name = f'Heading {level}'
            p = doc.add_paragraph(style=style_name)

            # For H1, we add inline runs (in case of bold/italic in heading)
            _add_inline_runs(p, heading_text)

            just_added_heading = True
            first_body_para_done = False  # next body para should be flush left
            i += 1
            continue

        # --- Horizontal rule ---
        if stripped in ('---', '***', '___'):
            flush_paragraph()
            close_lists()
            # Add a thin horizontal rule paragraph
            p = doc.add_paragraph(style='Normal')
            p.paragraph_format.first_line_indent = Pt(0)
            # Add a bottom border to simulate HR
            pPr = p._element.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                p._element.insert(0, pPr)
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '4')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'CCCCCC')
            pBdr.append(bottom)
            pPr.append(pBdr)
            just_added_heading = False
            i += 1
            continue

        # --- Bullet list items ---
        bullet_match = re.match(r'^[-*+]\s+(.+)$', stripped)
        if bullet_match:
            flush_paragraph()
            # If we were in a numbered list, close it
            if in_numbered_list:
                in_numbered_list = False

            item_text = bullet_match.group(1).strip()

            p = doc.add_paragraph(style='List Paragraph')
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.25)  # hanging indent
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)

            # Add bullet character
            bullet_run = p.add_run('\u2022  ')  # bullet + two spaces
            bullet_run.font.name = GARAMOND
            _set_run_fonts(bullet_run, GARAMOND)
            bullet_run.font.size = Pt(11)

            _add_inline_runs(p, item_text)

            in_bullet_list = True
            just_added_heading = False
            i += 1
            continue

        # --- Numbered list items ---
        num_match = re.match(r'^(\d+)\.\s+(.+)$', stripped)
        if num_match:
            flush_paragraph()
            # If we were in a bullet list, close it
            if in_bullet_list:
                in_bullet_list = False

            num_str = num_match.group(1)
            item_text = num_match.group(2).strip()

            p = doc.add_paragraph(style='List Paragraph')
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.25)  # hanging indent
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)

            # Add number
            num_run = p.add_run(f'{num_str}.  ')
            num_run.font.name = GARAMOND
            _set_run_fonts(num_run, GARAMOND)
            num_run.font.size = Pt(11)

            _add_inline_runs(p, item_text)

            in_numbered_list = True
            just_added_heading = False
            i += 1
            continue

        # --- Regular text line (accumulate into paragraph) ---
        pending_para_lines.append(stripped)
        # NOTE: Do NOT set just_added_heading = False here!
        # The flush_paragraph() function uses just_added_heading to determine
        # whether to suppress first-line indent. It clears the flag itself.
        i += 1

    # Flush any remaining content
    if in_blockquote:
        flush_blockquote()
    flush_paragraph()
    close_lists()

    return just_added_heading


# ============================================================================
# PAGE CONSTRUCTION
# ============================================================================

def _add_page_break(doc: Document) -> None:
    """Add a page break paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def _add_title_page(doc: Document, meta: dict) -> None:
    """Add the title page: Title (28pt, bold, centered), Subtitle (16pt, italic, gray), Author (14pt, gray)."""
    # Add some vertical space before title
    spacer = doc.add_paragraph()
    spacer.paragraph_format.first_line_indent = Pt(0)
    spacer.paragraph_format.space_before = Pt(120)
    spacer.paragraph_format.space_after = Pt(0)

    # Title
    title_p = doc.add_paragraph()
    title_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.first_line_indent = Pt(0)
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(12)
    title_run = title_p.add_run(meta.get('title', 'Untitled'))
    title_run.font.name = GARAMOND
    _set_run_fonts(title_run, GARAMOND)
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    # Subtitle
    if meta.get('subtitle'):
        sub_p = doc.add_paragraph()
        sub_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_p.paragraph_format.first_line_indent = Pt(0)
        sub_p.paragraph_format.space_before = Pt(0)
        sub_p.paragraph_format.space_after = Pt(24)
        sub_run = sub_p.add_run(meta['subtitle'])
        sub_run.font.name = GARAMOND
        _set_run_fonts(sub_run, GARAMOND)
        sub_run.font.size = Pt(16)
        sub_run.font.italic = True
        sub_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Author
    author_p = doc.add_paragraph()
    author_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_p.paragraph_format.first_line_indent = Pt(0)
    author_p.paragraph_format.space_before = Pt(36)
    author_p.paragraph_format.space_after = Pt(6)
    author_run = author_p.add_run(f"by {meta.get('author', 'Unknown')}")
    author_run.font.name = GARAMOND
    _set_run_fonts(author_run, GARAMOND)
    author_run.font.size = Pt(14)
    author_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Page break after title page
    _add_page_break(doc)


def _add_copyright_page(doc: Document, meta: dict) -> None:
    """Add the copyright page in 8pt Copyright style."""
    author = meta.get('author', 'Unknown')
    publisher = meta.get('publisher', author)
    year = meta.get('date', str(datetime.now().year))
    if len(year) > 4:
        year = year[:4]
    title = meta.get('title', 'Untitled')

    lines = [
        f"Copyright \u00a9 {year} by {author}",
        "",
        "All rights reserved. No part of this publication may be reproduced, distributed, "
        "or transmitted in any form or by any means, including photocopying, recording, "
        "or other electronic or mechanical methods, without the prior written permission "
        "of the publisher, except for brief quotations in reviews and certain "
        "noncommercial uses permitted by copyright law.",
        "",
        f"Published {year} by {publisher}",
        "",
        f"{' '.join(word[0].upper() + word[1:] for word in title.lower().split())}",
    ]

    for line in lines:
        p = doc.add_paragraph(style='Copyright')
        p.paragraph_format.first_line_indent = Pt(0)
        if line:
            run = p.add_run(line)
            run.font.name = GARAMOND
            _set_run_fonts(run, GARAMOND)

    # Page break after copyright
    _add_page_break(doc)


def _add_toc_page(doc: Document, chapters: list) -> None:
    """
    Add the Table of Contents page.
    Title "Contents" (20pt, centered) with a "toc" bookmark for KDP detection.
    Each entry is a hyperlink to its chapter bookmark, with dot leaders.
    """
    # TOC title
    toc_title = doc.add_paragraph()
    toc_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_title.paragraph_format.first_line_indent = Pt(0)
    toc_title.paragraph_format.space_before = Pt(24)
    toc_title.paragraph_format.space_after = Pt(18)
    toc_title.paragraph_format.page_break_before = True

    toc_run = toc_title.add_run('Contents')
    toc_run.font.name = GARAMOND
    _set_run_fonts(toc_run, GARAMOND)
    toc_run.font.size = Pt(20)
    toc_run.font.bold = True
    toc_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    # Add "toc" bookmark on the TOC title paragraph for KDP detection
    _add_bookmark(toc_title, 'toc')

    # TOC entries
    for ch in chapters:
        ch_num = ch.get('number', 0)
        ch_title = ch.get('title', f'Chapter {ch_num}')
        bookmark_name = f'chapter-{ch_num:02d}'

        entry_p = doc.add_paragraph()
        entry_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        entry_p.paragraph_format.first_line_indent = Pt(0)
        entry_p.paragraph_format.space_before = Pt(2)
        entry_p.paragraph_format.space_after = Pt(2)
        entry_p.paragraph_format.left_indent = Pt(0)
        entry_p.paragraph_format.right_indent = Pt(0)

        # Add dot leader tab stop
        _add_dot_leader_tab(entry_p, 6.0)

        # Add hyperlink to chapter bookmark
        _add_hyperlink(entry_p, bookmark_name, ch_title)

        # Add tab (dot leader) and page number placeholder
        # The page number is a placeholder since we can't compute real page numbers
        # until Word renders the document. We use a PAGE field for the bookmark.
        tab_run = entry_p.add_run('\t')
        tab_run.font.name = GARAMOND
        _set_run_fonts(tab_run, GARAMOND)
        tab_run.font.size = Pt(11)
        tab_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

        # Page number as a PAGREF field referencing the chapter bookmark
        # This shows as "?" until Word updates fields (Ctrl+A, F9)
        _add_pageref_field(entry_p, bookmark_name)

    # Page break after TOC
    _add_page_break(doc)


def _add_pageref_field(paragraph, bookmark_name: str) -> None:
    """
    Add a PAGREF field to a paragraph that references a bookmark.
    In Word, this displays the page number where the bookmark is located.
    It shows as "?" until fields are updated in Word.
    """
    run = OxmlElement('w:r')

    rpr = OxmlElement('w:rPr')
    rfonts = OxmlElement('w:rFonts')
    rfonts.set(qn('w:ascii'), GARAMOND)
    rfonts.set(qn('w:hAnsi'), GARAMOND)
    rfonts.set(qn('w:cs'), GARAMOND)
    rpr.append(rfonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '22')
    rpr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), '22')
    rpr.append(szCs)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '444444')
    rpr.append(color)
    run.append(rpr)

    # Field character: begin
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    run.append(fldChar_begin)

    paragraph._element.append(run)

    # Field code: PAGREF bookmark_name
    run2 = OxmlElement('w:r')
    rpr2 = OxmlElement('w:rPr')
    rfonts2 = OxmlElement('w:rFonts')
    rfonts2.set(qn('w:ascii'), GARAMOND)
    rfonts2.set(qn('w:hAnsi'), GARAMOND)
    rfonts2.set(qn('w:cs'), GARAMOND)
    rpr2.append(rfonts2)
    run2.append(rpr2)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = f' PAGREF {bookmark_name} '
    run2.append(instrText)
    paragraph._element.append(run2)

    # Field character: separate
    run3 = OxmlElement('w:r')
    fldChar_sep = OxmlElement('w:fldChar')
    fldChar_sep.set(qn('w:fldCharType'), 'separate')
    run3.append(fldChar_sep)
    paragraph._element.append(run3)

    # Display text (placeholder)
    run4 = OxmlElement('w:r')
    rpr4 = OxmlElement('w:rPr')
    rfonts4 = OxmlElement('w:rFonts')
    rfonts4.set(qn('w:ascii'), GARAMOND)
    rfonts4.set(qn('w:hAnsi'), GARAMOND)
    rfonts4.set(qn('w:cs'), GARAMOND)
    rpr4.append(rfonts4)
    sz4 = OxmlElement('w:sz')
    sz4.set(qn('w:val'), '22')
    rpr4.append(sz4)
    color4 = OxmlElement('w:color')
    color4.set(qn('w:val'), '444444')
    rpr4.append(color4)
    run4.append(rpr4)
    t4 = OxmlElement('w:t')
    t4.set(qn('xml:space'), 'preserve')
    t4.text = '?'  # Placeholder until Word updates fields
    run4.append(t4)
    paragraph._element.append(run4)

    # Field character: end
    run5 = OxmlElement('w:r')
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run5.append(fldChar_end)
    paragraph._element.append(run5)


def _add_chapter(doc: Document, chapter: dict) -> None:
    """
    Add a chapter to the document.
    - Chapter title as Heading 1 with page break before
    - Chapter bookmark (chapter-01 through chapter-12)
    - Chapter content parsed from markdown
    """
    ch_num = chapter.get('number', 0)
    ch_title = chapter.get('title', f'Chapter {ch_num}')
    content = chapter.get('content', '')

    # Strip ALL H1 lines from chapter content.
    # The chapter title is already added separately as Heading 1.
    # Any remaining H1 lines in the body are demoted to H2 (## ).
    # This prevents duplicate H1 headings that confuse KDP and Word navigation.
    content_lines = content.split('\n')
    cleaned_lines = []
    h1_stripped = False
    for line in content_lines:
        stripped_line = line.strip()
        # Match any H1 line: "# something" (but not "## something")
        if re.match(r'^#[^#]', stripped_line):
            if not h1_stripped:
                # Skip the first H1 (it's the chapter title, already added)
                h1_stripped = True
                continue
            else:
                # Demote subsequent H1s to H2
                demoted = '## ' + stripped_line[2:].strip()
                cleaned_lines.append(demoted)
                continue
        cleaned_lines.append(line)

    content = '\n'.join(cleaned_lines)

    # Add chapter heading (Heading 1 style has page_break_before)
    heading_p = doc.add_paragraph(style='Heading 1')
    _add_inline_runs(heading_p, ch_title)

    # Add chapter bookmark
    bookmark_name = f'chapter-{ch_num:02d}'
    _add_bookmark(heading_p, bookmark_name)

    # Parse chapter content (first para after heading = flush left)
    if content.strip():
        parse_markdown_to_docx(doc, content, first_para_after_heading=True)


def _add_about_author(doc: Document, meta: dict) -> None:
    """Add 'About the Author' section with Heading 2 and body text."""
    author = meta.get('author', 'Unknown')

    heading = doc.add_paragraph(style='Heading 2')
    heading.paragraph_format.page_break_before = True
    _add_inline_runs(heading, 'About the Author')

    bio_text = (
        f"{author} is the author of multiple practical guides on productivity, health, "
        f"technology, and personal development. His books are designed to deliver real, "
        f"actionable information — no fluff, no filler, just results."
    )

    p1 = doc.add_paragraph(style='Normal')
    p1.paragraph_format.first_line_indent = Pt(0)  # flush left after heading
    _add_inline_runs(p1, bio_text)

    p2 = doc.add_paragraph(style='Normal')
    p2.paragraph_format.first_line_indent = Inches(0.25)
    _add_inline_runs(p2, f'For more books and resources, search "{author}" on Amazon.')


def _add_cta(doc: Document, meta: dict) -> None:
    """Add call-to-action back matter: centered text with URL."""
    author = meta.get('author', 'Unknown')
    author_slug = author.lower().replace(' ', '')

    heading = doc.add_paragraph(style='Heading 2')
    heading.paragraph_format.page_break_before = True
    _add_inline_runs(heading, f'Also by {author}')

    cta_text = doc.add_paragraph()
    cta_text.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cta_text.paragraph_format.first_line_indent = Pt(0)
    cta_text.paragraph_format.space_before = Pt(12)
    cta_text.paragraph_format.space_after = Pt(6)

    run1 = cta_text.add_run('Find all books at: ')
    run1.font.name = GARAMOND
    _set_run_fonts(run1, GARAMOND)
    run1.font.size = Pt(11)

    run2 = cta_text.add_run(f'amazon.com/author/{author_slug}')
    run2.font.name = GARAMOND
    _set_run_fonts(run2, GARAMOND)
    run2.font.size = Pt(11)
    run2.font.bold = True


# ============================================================================
# MAIN BUILDER
# ============================================================================

def build_docx(meta: dict, chapters: list, output_dir: Path) -> Path:
    """
    Build a complete DOCX file from metadata and chapter content.

    Args:
        meta: dict with keys: title, subtitle, author, slug, date, language
              (also accepts: publisher, description, keywords, rights)
        chapters: list of dicts with keys: number, title, content (markdown string)
        output_dir: Path to the directory where the .docx file will be saved

    Returns:
        Path to the generated .docx file
    """
    global _next_bookmark_id
    _next_bookmark_id = 100  # reset for each build

    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    slug = meta.get('slug', meta.get('title', 'book').lower().replace(' ', '-'))
    docx_path = output_dir / f'{slug}.docx'

    log(f'Building DOCX: {docx_path.name}')

    # Create document
    doc = Document()

    # Set default document font
    style = doc.styles['Normal']
    font = style.font
    font.name = GARAMOND
    font.size = Pt(11)

    # Setup all custom styles
    _setup_styles(doc)

    # --- Page order ---

    # 1. Title page
    log('  Adding title page')
    _add_title_page(doc, meta)

    # 2. Copyright page
    log('  Adding copyright page')
    _add_copyright_page(doc, meta)

    # 3. TOC page
    log('  Adding table of contents')
    _add_toc_page(doc, chapters)

    # 4. Chapters 1-12
    for ch in chapters:
        ch_num = ch.get('number', 0)
        ch_title = ch.get('title', f'Chapter {ch_num}')
        words = len(ch.get('content', '').split())
        log(f'  Adding Chapter {ch_num:02d}: {ch_title} ({words:,} words)')
        _add_chapter(doc, ch)

    # 5. About the Author
    log('  Adding About the Author')
    _add_about_author(doc, meta)

    # 6. CTA
    log('  Adding CTA back matter')
    _add_cta(doc, meta)

    # Save
    doc.save(str(docx_path))
    log(f'DOCX saved: {docx_path} ({docx_path.stat().st_size:,} bytes)')

    return docx_path


# ============================================================================
# CLI (standalone execution)
# ============================================================================

def load_polished_chapters(workbook_dir: Path) -> list:
    """
    Load all approved chapters from w-polished/ directory.
    Returns list of dicts sorted by chapter number.
    """
    polished_dir = workbook_dir / 'w-polished'

    if not polished_dir.exists():
        print(f'ERROR: w-polished/ directory not found: {polished_dir}', file=sys.stderr)
        sys.exit(1)

    chapter_files = sorted(polished_dir.glob('chapter-*.md'))

    if not chapter_files:
        print(f'ERROR: No approved chapters found in {polished_dir}', file=sys.stderr)
        sys.exit(1)

    chapters = []
    for path in chapter_files:
        content = path.read_text(encoding='utf-8')

        # Extract chapter number from filename
        num_match = re.search(r'chapter-(\d+)\.md', path.name)
        num = int(num_match.group(1)) if num_match else 0

        # Extract title from first H1
        title_match = re.search(r'^#\s+Chapter\s+\d+:\s*(.+)', content, re.M)
        if not title_match:
            title_match = re.search(r'^#\s+(.+)', content, re.M)
        title = title_match.group(1).strip() if title_match else f'Chapter {num}'

        chapters.append({
            'number': num,
            'title': title,
            'content': content,
        })

    chapters.sort(key=lambda c: c['number'])
    log(f'Loaded {len(chapters)} approved chapters')
    return chapters


def extract_book_metadata(workbook_dir: Path) -> dict:
    """Extract basic metadata from the workbook directory."""
    outline_path = workbook_dir / '01_outline.md'
    meta = {
        'title': workbook_dir.name.replace('book-', '').replace('-', ' ').title(),
        'subtitle': '',
        'author': 'William Archer',
        'slug': workbook_dir.name.replace('book-', ''),
        'date': str(datetime.now().year),
        'language': 'en',
        'publisher': 'William Archer',
    }

    if outline_path.exists():
        content = outline_path.read_text(encoding='utf-8')
        # Try to extract title from outline
        title_match = re.search(r'^#\s+(.+)', content, re.M)
        if title_match:
            raw_title = title_match.group(1).strip()
            # Strip common prefixes from auto-generated outlines
            raw_title = re.sub(r'^Book\s+Outline:\s*', '', raw_title, flags=re.IGNORECASE)
            meta['title'] = raw_title

        # Try to extract subtitle
        sub_match = re.search(r'^##\s+Subtitle:\s*(.+)', content, re.M)
        if sub_match:
            meta['subtitle'] = sub_match.group(1).strip()

    return meta


def main():
    """CLI entry point for standalone execution."""
    import argparse

    parser = argparse.ArgumentParser(
        description='DOCX Builder — python-docx native generation for Ebook Factory',
    )
    parser.add_argument('--book-dir', type=str, default=None,
                        help='Workbook directory (auto-detects latest if not set)')
    parser.add_argument('--author', type=str, default=None,
                        help='Author name override')
    parser.add_argument('--title', type=str, default=None,
                        help='Title override')

    args = parser.parse_args()

    # Find workbook directory
    workbook_dir = None
    if args.book_dir:
        workbook_dir = Path(args.book_dir)
    else:
        # Auto-detect latest workbook
        hermes_home = Path(os.environ.get('HERMES_HOME', Path.home() / '.hermes'))
        workbooks_dir = hermes_home / 'ebook-factory' / 'workbooks'
        if workbooks_dir.exists():
            workbooks = sorted(workbooks_dir.glob('book-*'))
            if workbooks:
                workbook_dir = workbooks[-1]

    if not workbook_dir or not workbook_dir.exists():
        print('ERROR: No workbook directory found. Use --book-dir or ensure w-polished/ exists.')
        sys.exit(1)

    log(f'Workbook: {workbook_dir}')

    # Load metadata
    meta = extract_book_metadata(workbook_dir)
    if args.author:
        meta['author'] = args.author
    if args.title:
        meta['title'] = args.title
        meta['slug'] = args.title.lower().replace(' ', '-')

    # Load chapters
    chapters = load_polished_chapters(workbook_dir)
    for ch in chapters:
        words = len(ch['content'].split())
        log(f'  Chapter {ch["number"]:02d}: {ch["title"]} ({words:,} words)')

    # Build DOCX
    output_dir = workbook_dir / 'output'
    docx_path = build_docx(meta, chapters, output_dir)

    log(f'\nDOCX build complete: {docx_path}')
    log(f'File size: {docx_path.stat().st_size:,} bytes')


if __name__ == '__main__':
    main()
