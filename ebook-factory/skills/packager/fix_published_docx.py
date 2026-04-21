#!/usr/bin/env python3
"""Batch-fix published-book DOCX files with title page + TOC + chapter formatting.

These books were built before the post-processor was added. They have:
- Title/Subtitle/Author/Date styles from pandoc --metadata
- No centered title page formatting
- Tiny TOC entries
- No chapter page breaks or bookmarks

Usage:
  python3 fix_published_docx.py
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from pathlib import Path
import re
import shutil

WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def fix_published_docx(docx_path):
    """Apply title page + TOC + chapter formatting to published-book DOCX."""
    # Backup
    bak = docx_path.with_suffix('.docx.bak')
    if not bak.exists():
        shutil.copy(str(docx_path), str(bak))

    doc = Document(str(docx_path))
    paragraphs = doc.paragraphs

    title_text = paragraphs[0].text.strip() if paragraphs else ''

    # Phase 1: Style the title page (Title + Subtitle + Author at top)
    for i, p in enumerate(paragraphs[:5]):
        if p.style.name == 'Title' and i == 0:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(28)
                run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
            pf = p.paragraph_format
            pf.space_before = Pt(72)
            pf.space_after = Pt(12)
            print('  Styled title')

        elif p.style.name == 'Subtitle':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(16)
                run.font.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            p.paragraph_format.space_after = Pt(24)
            print('  Styled subtitle')

        elif p.style.name == 'Author':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            p.paragraph_format.space_before = Pt(36)
            p.paragraph_format.space_after = Pt(6)
            # Add page break after author
            run_elem = p._element.makeelement(qn('w:br'), {qn('w:type'): 'page'})
            new_run = p.add_run()
            new_run._element.append(run_elem)
            print('  Styled author + page break')

        elif p.style.name == 'Date':
            # Remove the Date line (redundant with copyright)
            parent = p._element.getparent()
            parent.remove(p._element)
            print('  Removed Date line')

    # Phase 2: Find and style TOC
    in_toc = False
    toc_count = 0
    paragraphs = doc.paragraphs  # re-read after removals

    for i, p in enumerate(paragraphs):
        text = p.text.strip()

        # Detect TOC heading
        if 'Table of Contents' in text and ('Heading' in p.style.name or p.style.name == 'Heading 1'):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(20)
                run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(18)
            in_toc = True
            toc_count = 0
            print('  Styled TOC heading (centered, 20pt)')
            continue

        # Style TOC entries (Compact or Normal style within TOC)
        if in_toc and (p.style.name == 'Compact' or (p.style.name == 'Normal' and 'Chapter' in text)):
            toc_count += 1
            # Style via XML (handles hyperlinks)
            for r_elem in p._element.iter('{' + WNS + '}r'):
                rpr = r_elem.find('{' + WNS + '}rPr')
                if rpr is None:
                    rpr = r_elem.makeelement('{' + WNS + '}rPr', {})
                    r_elem.insert(0, rpr)
                sz = rpr.find('{' + WNS + '}sz')
                if sz is None:
                    sz = r_elem.makeelement('{' + WNS + '}sz', {})
                    rpr.append(sz)
                sz.set('{' + WNS + '}val', '26')
                szCs = rpr.find('{' + WNS + '}szCs')
                if szCs is None:
                    szCs = r_elem.makeelement('{' + WNS + '}szCs', {})
                    rpr.append(szCs)
                szCs.set('{' + WNS + '}val', '26')
                color = rpr.find('{' + WNS + '}color')
                if color is None:
                    color = r_elem.makeelement('{' + WNS + '}color', {})
                    rpr.append(color)
                color.set('{' + WNS + '}val', '333333')
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.space_before = Pt(2)

            # End TOC on non-TOC style
            if p.style.name == 'Normal':
                if toc_count >= 8 or (i + 1 < len(paragraphs) and paragraphs[i+1].style.name not in ('Compact', 'Normal')):
                    in_toc = False
                    run_elem = p._element.makeelement(qn('w:br'), {qn('w:type'): 'page'})
                    new_run = p.add_run()
                    new_run._element.append(run_elem)
                    print('  Styled %d TOC entries + page break after TOC' % toc_count)
            continue

        if in_toc and p.style.name not in ('Compact', 'Normal'):
            in_toc = False
            # Add page break on the last TOC entry
            if i > 0:
                prev = paragraphs[i-1]
                run_elem = prev._element.makeelement(qn('w:br'), {qn('w:type'): 'page'})
                new_run = prev.add_run()
                new_run._element.append(run_elem)
                print('  Page break after TOC')

        # Phase 3: Add page breaks before chapter headings
        if 'Chapter' in text and 'Heading' in p.style.name:
            p.paragraph_format.page_break_before = True
            # Add bookmark if not already present
            ns = {'w': WNS}
            existing_bms = p._element.findall('.//w:bookmarkStart', ns)
            if not existing_bms:
                ch_match = re.search(r'Chapter\s+(\d+)', text)
                if ch_match:
                    ch_num = int(ch_match.group(1))
                    bm_name = 'chapter-%02d' % ch_num
                    bm_id = 100 + ch_num
                    bm_start = p._element.makeelement('{' + WNS + '}bookmarkStart', {
                        '{' + WNS + '}id': str(bm_id),
                        '{' + WNS + '}name': bm_name,
                    })
                    p._element.insert(0, bm_start)
                    bm_end = p._element.makeelement('{' + WNS + '}bookmarkEnd', {
                        '{' + WNS + '}id': str(bm_id),
                    })
                    p._element.append(bm_end)
                    print('  Page break + bookmark %s' % bm_name)
                else:
                    print('  Page break before: %s' % text[:50])

    doc.save(str(docx_path))
    print('  Saved: ' + str(docx_path))


if __name__ == '__main__':
    pub_dir = Path('/home/bookforge/books/factory/references/published-books')
    for d in sorted(pub_dir.iterdir()):
        if not d.is_dir():
            continue
        docx_files = list(d.glob('*.docx'))
        if not docx_files:
            continue
        print('=== ' + d.name + ' ===')
        try:
            fix_published_docx(docx_files[0])
        except Exception as e:
            print('  ERROR: ' + str(e))
        print()
