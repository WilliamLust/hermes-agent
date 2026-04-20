#!/usr/bin/env python3
"""
Packager Agent — Ebook Factory Phase 3
=======================================

Assembles all approved chapters from w-polished/ into a complete ebook package:
  - Full HTML manuscript
  - EPUB (via ebooklib)
  - PDF (via WeasyPrint, or fallback Calibre ebook-convert)
  - KDP metadata JSON
  - Front/back matter (title page, copyright, TOC, author bio, CTA)

Only runs after human review has approved chapters into w-polished/.

Usage:
    python3 packager.py
    python3 packager.py --book-dir ~/.hermes/ebook-factory/workbooks/book-my-topic/
    python3 packager.py --author "William Archer" --title "My Book Title"
    python3 packager.py --formats epub,pdf  # Comma-separated: epub, pdf, html

Reference: ~/hermes-agent/AGENTS.md section 6
"""

import os
import sys
import re
import json
import shutil
import argparse
import subprocess
import requests
import html
from pathlib import Path
from datetime import datetime

# ============================================================================
# CONFIGURATION
# ============================================================================

HERMES_HOME = Path(os.environ.get("HERMES_HOME", Path.home() / ".hermes"))
WORKBOOKS_DIR = HERMES_HOME / "ebook-factory" / "workbooks"
STYLE_GUIDE = Path.home() / "books" / "factory" / "style-guide.md"

DEFAULT_AUTHOR = "William Archer"
DEFAULT_LANGUAGE = "en"

OLLAMA_URL     = os.environ.get("OLLAMA_BASE_URL", "http://localhost:11434") + "/api/chat"
METADATA_MODEL = "qwen3.5:27b-16k"   # Description + subtitle generation


# ============================================================================
# LOGGING
# ============================================================================

def log(msg: str) -> None:
    ts = datetime.now().strftime("%H:%M:%S")
    print(f"[{ts}] {msg}", flush=True)


def log_section(title: str) -> None:
    print(f"\n{'='*60}", flush=True)
    print(f"  {title}", flush=True)
    print(f"{'='*60}", flush=True)


def error_exit(msg: str) -> None:
    print(f"\nERROR: {msg}", file=sys.stderr, flush=True)
    sys.exit(1)


def _fix_docx_title_page(docx_path: Path, meta: dict) -> None:
    """Post-process pandoc-generated DOCX to fix the title page formatting.

    Pandoc converts the HTML title-page div into bare paragraphs with no
    special styling. This function:
      1. Identifies the title-page paragraphs (h1.book-title + p.book-author etc.)
      2. Centers them and applies proper font sizes
      3. Adds page breaks after title and copyright pages
      4. Removes any pandoc auto-generated Title/Author paragraphs that
         duplicate the HTML title page
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        log("  python-docx not installed -- skipping DOCX title page fix")
        return

    doc = Document(str(docx_path))
    paragraphs = doc.paragraphs

    # Phase 1: Remove duplicate pandoc-generated Title/Author paragraphs
    # Pandoc creates "Title" and "Author" style paragraphs from metadata.
    # These duplicate the HTML title page. Remove them.
    to_remove = []
    for i, p in enumerate(paragraphs):
        if p.style.name in ("Title", "Author") and i < 5:
            to_remove.append(i)
            log(f"  Removing duplicate pandoc paragraph: style={p.style.name}, text={p.text[:40]}")

    # Remove in reverse order to preserve indices
    for i in reversed(to_remove):
        p = paragraphs[i]
        parent = p._element.getparent()
        parent.remove(p._element)

    # Re-read paragraphs after removal
    paragraphs = doc.paragraphs

    # Phase 2: Format the title page elements
    # The HTML title-page div produces these paragraphs in order:
    #   - h1.book-title  -> "Heading 1" style (the book title)
    #   - p.book-subtitle -> "First Paragraph" or body (the subtitle)
    #   - p.book-author  -> body text ("by Author Name")
    #   - p.book-publisher -> body text (publisher name)
    #   - p.book-date    -> body text (year)
    title_text = meta.get("title", "")
    author_text = f"by {meta.get('author', 'William Archer')}"
    subtitle_text = meta.get("subtitle", "")
    year_text = str(meta.get("date", ""))[:4] if meta.get("date") else ""

    in_title_page = False
    title_page_done = False
    copyright_done = False

    for i, p in enumerate(paragraphs):
        text = p.text.strip()

        # Detect start of title page (the Heading 1 with the book title)
        if not title_page_done and p.style.name == "Heading 1" and text == title_text:
            in_title_page = True
            # Style the title: large, centered, bold
            p.style = doc.styles['Title']
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(28)
                run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
            # Add space before
            pf = p.paragraph_format
            pf.space_before = Pt(72)
            pf.space_after = Pt(12)
            log(f"  Styled title: '{text[:50]}'")
            continue

        if in_title_page and not title_page_done:
            # Subtitle
            if subtitle_text and text == subtitle_text:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(16)
                    run.font.italic = True
                    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                pf = p.paragraph_format
                pf.space_after = Pt(24)
                log(f"  Styled subtitle: '{text[:50]}'")
                continue

            # Author line
            if text == author_text or (text.startswith("by ") and meta.get("author", "") in text):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(14)
                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                pf = p.paragraph_format
                pf.space_before = Pt(36)
                pf.space_after = Pt(6)
                log(f"  Styled author: '{text}'")
                continue

            # Publisher name (matches author name)
            if text == meta.get("publisher", ""):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
                continue

            # Year
            if text == year_text or (len(text) == 4 and text.isdigit()):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
                # Page break after title page
                pf = p.paragraph_format
                pf.space_after = Pt(0)
                from docx.oxml.ns import qn
                run_elem = p._element.makeelement(qn('w:br'), {qn('w:type'): 'page'})
                # Add page break as a new run at end
                new_run = p.add_run()
                new_run._element.append(run_elem)
                title_page_done = True
                in_title_page = False
                log(f"  Styled year + page break after title page")
                continue

        # Add page break after copyright page
        if not copyright_done and "All rights reserved" in text and "reproduced" in text:
            # Find the last paragraph of the copyright section
            # (it ends with "Published YYYY by...")
            pass
        if not copyright_done and text.startswith("Published") and meta.get("author", "") in text:
            from docx.oxml.ns import qn
            run_elem = p._element.makeelement(qn('w:br'), {qn('w:type'): 'page'})
            new_run = p.add_run()
            new_run._element.append(run_elem)
            copyright_done = True
            log(f"  Added page break after copyright page")
            continue

        # Add page break before each chapter heading (Chapter N: ...)
        if p.style.name == "Heading 1" and "Chapter" in text:
            p.paragraph_format.page_break_before = True
            log(f"  Page break before: '{text[:60]}'")
            continue

    doc.save(str(docx_path))
    log(f"  DOCX title page formatting applied")


# ============================================================================
# FILE I/O
# ============================================================================

def read_file(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return path.read_text(encoding="latin-1")


def write_file_atomic(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    tmp = path.with_suffix(".tmp")
    tmp.write_text(content, encoding="utf-8")
    shutil.move(str(tmp), str(path))
    log(f"Written: {path}")


def find_latest_workbook() -> Path:
    if not WORKBOOKS_DIR.exists():
        error_exit(f"Workbooks directory not found: {WORKBOOKS_DIR}")
    books = [d for d in WORKBOOKS_DIR.iterdir() if d.is_dir() and d.name.startswith("book-")]
    if not books:
        error_exit("No workbooks found. Run outliner + chapter-builder first.")
    return sorted(books, key=lambda d: d.stat().st_mtime, reverse=True)[0]


# ============================================================================
# LLM METADATA GENERATION
# ============================================================================

# Quality bar — the Procrastination Fix description format is what we're matching:
# hook paragraph → what the book reveals (named bullet features) → reader callout
_DESCRIPTION_EXAMPLE = """\
**You're not lazy. You're stuck. And there's a difference.**

You know exactly what you should be doing. That project that matters. That deadline looming. \
But here you are, doing literally anything except the thing that actually matters.

**The Procrastination Fix** reveals what research has known for years: procrastination isn't \
about laziness. It's an emotional regulation problem. You're not avoiding the task — you're \
avoiding the anxiety, boredom, or self-doubt it triggers.

This book gives you the tools to break that pattern:

• **The 2-Minute Bridge** — A micro-commitment technique that tricks your brain past the starting line
• **Emotion-First Planning** — Stop planning around tasks. Start planning around feelings.
• **The Zero-Draft Method** — Why permission to produce garbage is the fastest path to quality work
• **The 80% Rule** — When good enough beats perfect every single time

No motivational fluff. No generic advice. Just specific, research-backed techniques that work \
even when you don't feel like it.

Whether you're a perfectionist who freezes, a deadline-driven crammer, or someone who works \
more hours than everyone else but on all the wrong things — this book meets you where you are.

Ready to stop stalling and start doing? Your future self is waiting.\
"""

_SUBTITLE_EXAMPLE = "Evidence-Based Strategies to Stop Stalling and Start Doing"

# KDP HTML description example (same content, formatted for KDP's CKEditor)
_HTML_DESCRIPTION_EXAMPLE = """\
<p><strong>You're not lazy. You're stuck. And there's a difference.</strong></p>
<p>You know exactly what you should be doing. That project that matters. That deadline looming. \
But here you are, doing literally anything except the thing that actually matters.</p>
<p><strong>The Procrastination Fix</strong> reveals what research has known for years: \
procrastination isn't about laziness. It's an emotional regulation problem.</p>
<p>This book gives you the tools to break that pattern:</p>
<ul>
<li><strong>The 2-Minute Bridge</strong> — A micro-commitment technique that tricks your brain past the starting line</li>
<li><strong>Emotion-First Planning</strong> — Stop planning around tasks. Start planning around feelings.</li>
<li><strong>The Zero-Draft Method</strong> — Why permission to produce garbage is the fastest path to quality work</li>
<li><strong>The 80% Rule</strong> — When good enough beats perfect every single time</li>
</ul>
<p>No motivational fluff. No generic advice. Just specific, research-backed techniques that work \
even when you don't feel like it.</p>
<p>Ready to stop stalling and start doing? Your future self is waiting.</p>\
"""


def generate_llm_metadata(title: str, outline_content: str) -> dict:
    """
    Use Qwen 27B to generate full marketing copy:
    - subtitle (5-9 words, specific promise)
    - description_plain (multi-paragraph with bullet features, ~200-280 words)
    - description_html (same content formatted for KDP's CKEditor)
    - gumroad_description (slightly longer, Gumroad-optimized)
    - receipt_message (personal thank-you + quick start tips)
    - keywords (7 specific Amazon search phrases)

    Returns dict with all fields, or empty strings on failure.
    Falls back silently — packager continues with placeholders if LLM fails.
    """
    ch_titles = re.findall(r"^## Chapter \d+:\s*(.+)", outline_content, re.MULTILINE)
    ch_summary = "\n".join(f"- {t}" for t in ch_titles[:12]) if ch_titles else "(chapters not found)"

    promise_match = re.search(r"\*\*Core Promise:\*\*\s*(.+)", outline_content)
    promise = promise_match.group(1).strip() if promise_match else ""
    reader_match = re.search(r"\*\*Target Reader:\*\*\s*(.+)", outline_content)
    reader = reader_match.group(1).strip() if reader_match else ""

    system = """You are a professional nonfiction book marketer who writes Amazon KDP copy that converts browsers to buyers.

Your descriptions follow this exact structure:
1. Bold hook sentence that names the reader's exact pain
2. 1-2 sentences validating their frustration with existing solutions
3. Bold book title + what it reveals (1 sentence)
4. 4-6 bullet features, each formatted as: • **Feature Name** — specific benefit or technique
5. 1-2 sentences: no fluff, just what this book delivers
6. 1-2 sentences: who this is for (specific, not generic)
7. Closing call to action (1 short line)

Rules:
- Write in second person (you/your throughout)
- Every bullet names a specific technique or framework from the book
- No hollow phrases: no "comprehensive", "journey", "whether you're a beginner or expert"
- Specific research/data beats vague claims
- HTML version uses <p>, <strong>, <ul>, <li> only (KDP supports these)
- Gumroad description is the same content with → bullet points instead of •
- Receipt message is personal, warm, gives 2-3 quick-start tips, signed "William Archer"
- Keywords are 2-4 word Amazon search phrases buyers actually type

Output JSON only with these exact keys:
{
  "subtitle": "...",
  "description_plain": "...",
  "description_html": "...",
  "gumroad_description": "...",
  "receipt_message": "...",
  "keywords": ["...", "...", "...", "...", "...", "...", "..."]
}"""

    user = f"""Write complete Amazon KDP + Gumroad marketing copy for this nonfiction ebook.

TITLE: {title}
CORE PROMISE: {promise}
TARGET READER: {reader}

CHAPTERS (use these to write specific bullet features):
{ch_summary}

QUALITY STANDARD — Match this format exactly (The Procrastination Fix):

Plain description:
---
{_DESCRIPTION_EXAMPLE}
---

HTML description:
---
{_HTML_DESCRIPTION_EXAMPLE}
---

Now write all marketing copy for "{title}".
Output ONLY valid JSON with all 6 keys."""

    log("Generating full marketing copy via Qwen 27B...")
    # Use shared ollama_client with retry
    sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
    from ollama_client import ollama_call_with_retry

    raw = ollama_call_with_retry(
        user, system, METADATA_MODEL,
        max_retries=2,
        num_predict=1800,
        temperature=0.7,
        timeout=300,
    )
    if raw is None:
        log("WARNING: LLM did not return metadata — using fallback")
        return {}
    raw = re.sub(r"<think/>.*?</think/>", "", raw, flags=re.DOTALL).strip()
    json_match = re.search(r"\{.*\}", raw, re.DOTALL)
    if not json_match:
        log("WARNING: LLM did not return JSON for metadata — using fallback")
        return {}
    try:
        data = json.loads(json_match.group(0))
    except json.JSONDecodeError:
        log("WARNING: LLM returned invalid JSON for metadata — using fallback")
        return {}
    # Validate all keys present
    required = ["subtitle", "description_plain", "description_html",
                "gumroad_description", "receipt_message", "keywords"]
    for key in required:
        if key not in data:
            data[key] = ""
    log(f"  Subtitle: {data.get('subtitle', '')}")
    log(f"  Description: {str(data.get('description_plain',''))[:80]}...")
    log(f"  Keywords: {data.get('keywords', [])}")
    return data


# ============================================================================
# METADATA EXTRACTION
# ============================================================================

def extract_book_metadata(workbook_dir: Path, outline_path: Path, cli_args) -> dict:
    """
    Extract book metadata from outline and/or CLI args.
    Falls back to sensible defaults from the workbook dir name.
    """
    outline_content = read_file(outline_path) if outline_path.exists() else ""

    # Try to extract title from outline
    title = None
    title_match = re.search(r'^#\s+Book\s+Outline:\s*(.+)', outline_content, re.M)
    if title_match:
        title = title_match.group(1).strip()

    # CLI arg overrides
    if cli_args.title:
        title = cli_args.title

    if not title:
        # Derive from workbook dir name: "book-my-topic" -> "My Topic"
        slug = workbook_dir.name.replace("book-", "").replace("-", " ")
        title = slug.title()

    author = cli_args.author if cli_args.author else DEFAULT_AUTHOR
    slug = workbook_dir.name

    # Load keywords from existing kdp-metadata.json if available (clean Qwen-generated phrases)
    # Fall back to scraping outline keywords only if no metadata file exists yet
    keywords = []
    existing_meta_path = workbook_dir / "output" / "kdp-metadata.json"
    if existing_meta_path.exists():
        try:
            existing_meta = json.loads(existing_meta_path.read_text(encoding="utf-8"))
            keywords = existing_meta.get("keywords", [])
            if keywords:
                log(f"Loaded {len(keywords)} keywords from existing kdp-metadata.json")
        except Exception:
            pass

    if not keywords:
        # Scrape from outline as fallback — strip markdown artifacts
        kw_match = re.search(r'[Kk]eywords?[^:]*:\s*(.+)', outline_content)
        if kw_match:
            raw = kw_match.group(1).strip().strip("[]")
            raw = re.sub(r'\*+', '', raw)  # strip bold markers
            keywords = [k.strip().strip("'\"") for k in raw.split(",") if k.strip() and len(k.strip()) > 2][:7]

    # Extract description from outline intro
    desc_match = re.search(r'(?:description|summary|about)[^:]*:\s*(.+?)(?:\n\n|\Z)', outline_content, re.I | re.DOTALL)
    description = desc_match.group(1).strip()[:200] if desc_match else f"A practical guide to {title}."

    # Generate LLM subtitle + description (overwrites regex fallback if successful)
    llm_meta   = generate_llm_metadata(title, outline_content)
    subtitle   = llm_meta.get("subtitle", "")
    if llm_meta.get("description_plain"):
        description = llm_meta["description_plain"]

    return {
        "title":               title,
        "subtitle":            subtitle,
        "author":              author,
        "slug":                slug,
        "language":            DEFAULT_LANGUAGE,
        "description":         description,
        "description_html":    llm_meta.get("description_html", ""),
        "gumroad_description": llm_meta.get("gumroad_description", ""),
        "receipt_message":     llm_meta.get("receipt_message", ""),
        "keywords":            keywords or llm_meta.get("keywords", []),
        "date":                datetime.now().strftime("%Y-%m-%d"),
        "publisher":           "William Archer",
        "rights":              f"Copyright {datetime.now().year} {author}. All rights reserved.",
    }


# ============================================================================
# FRONT / BACK MATTER
# ============================================================================

def generate_title_page(meta: dict) -> str:
    subtitle_html = f'\n<p class="book-subtitle">{meta["subtitle"]}</p>' if meta.get("subtitle") else ""
    return f"""<div class="title-page">
<h1 class="book-title">{meta['title']}</h1>{subtitle_html}
<p class="book-author">by {meta['author']}</p>
<p class="book-publisher">{meta['publisher']}</p>
<p class="book-date">{datetime.now().strftime('%Y')}</p>
</div>"""


def generate_copyright_page(meta: dict) -> str:
    return f"""<div class="copyright-page">
<p>{meta['rights']}</p>
<p>All rights reserved. No part of this publication may be reproduced, distributed,
or transmitted in any form or by any means, including photocopying, recording,
or other electronic or mechanical methods, without the prior written permission
of the publisher.</p>
<p>Published {meta['date']} by {meta['publisher']}</p>
</div>"""


def generate_author_bio(author: str) -> str:
    return f"""<div class="author-bio">
<h2>About the Author</h2>
<p>{author} is the author of multiple practical guides on productivity, health,
technology, and personal development. His books are designed to deliver real,
actionable information — no fluff, no filler, just results.</p>
<p>For more books and resources, search "{author}" on Amazon.</p>
</div>"""


def generate_cta(meta: dict, published_books: list) -> str:
    """Generate back-matter call-to-action listing other books."""
    books_html = ""
    for book in published_books[:5]:  # List up to 5 other books
        if book != meta["title"]:
            books_html += f'<li>{book}</li>\n'

    if not books_html:
        books_html = '<li>Search for more books by this author on Amazon</li>\n'

    return f"""<div class="also-by">
<h2>Also by {meta['author']}</h2>
<ul>
{books_html}
</ul>
<p>Find all books at: <strong>amazon.com/author/{meta['author'].lower().replace(' ', '')}</strong></p>
</div>"""


def generate_toc_html(chapters: list) -> str:
    """Generate HTML table of contents from chapter list."""
    items = ""
    for ch in chapters:
        items += f'<li><a href="#chapter-{ch["number"]:02d}">{ch["title"]}</a></li>\n'
    return f"""<div class="toc">
<h2>Table of Contents</h2>
<ol>
{items}
</ol>
</div>"""


# ============================================================================
# MARKDOWN TO HTML
# ============================================================================

def md_to_html(md_content: str, chapter_id: str = "") -> str:
    """
    Simple markdown-to-HTML converter.
    Handles: headings, bold, italic, lists, paragraphs, horizontal rules.
    Uses Python stdlib only — no external markdown library required.
    """
    # Strip validation report section (not for readers)
    md_content = re.sub(r'\n## Validation Report.*', '', md_content, flags=re.DOTALL)
    # Strip AUTO-GENERATED comment
    md_content = re.sub(r'<!--.*?-->', '', md_content, flags=re.DOTALL)

    lines = md_content.split("\n")
    html_lines = []
    in_list = False
    in_paragraph = False
    paragraph_lines = []

    def flush_paragraph():
        nonlocal in_paragraph, paragraph_lines
        if paragraph_lines:
            text = " ".join(paragraph_lines).strip()
            if text:
                html_lines.append(f"<p>{text}</p>")
        in_paragraph = False
        paragraph_lines = []

    def inline_format(text: str) -> str:
        # Bold
        text = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', text)
        # Italic
        text = re.sub(r'\*(.+?)\*', r'<em>\1</em>', text)
        # Code
        text = re.sub(r'`(.+?)`', r'<code>\1</code>', text)
        return text

    for line in lines:
        stripped = line.strip()

        # Close list if needed
        if in_list and not stripped.startswith(("- ", "* ", "+ ")) and not re.match(r'^\d+\.', stripped):
            html_lines.append("</ul>")
            in_list = False

        # Headings
        if stripped.startswith("# "):
            flush_paragraph()
            text = stripped[2:].strip()
            html_lines.append(f'<h1 id="{chapter_id}">{inline_format(text)}</h1>')
        elif stripped.startswith("## "):
            flush_paragraph()
            text = stripped[3:].strip()
            anchor = re.sub(r'[^a-z0-9]+', '-', text.lower())
            html_lines.append(f'<h2 id="{anchor}">{inline_format(text)}</h2>')
        elif stripped.startswith("### "):
            flush_paragraph()
            text = stripped[4:].strip()
            html_lines.append(f'<h3>{inline_format(text)}</h3>')
        elif stripped.startswith("#### "):
            flush_paragraph()
            text = stripped[5:].strip()
            html_lines.append(f'<h4>{inline_format(text)}</h4>')

        # Horizontal rule
        elif stripped in ("---", "***", "___"):
            flush_paragraph()
            html_lines.append("<hr/>")

        # Unordered list
        elif re.match(r'^[-*+]\s', stripped):
            if in_paragraph:
                flush_paragraph()
            if not in_list:
                html_lines.append("<ul>")
                in_list = True
            text = re.sub(r'^[-*+]\s', '', stripped)
            html_lines.append(f"<li>{inline_format(text)}</li>")

        # Ordered list
        elif re.match(r'^\d+\.\s', stripped):
            if in_paragraph:
                flush_paragraph()
            if not in_list:
                html_lines.append("<ol>")
                in_list = True
            text = re.sub(r'^\d+\.\s', '', stripped)
            html_lines.append(f"<li>{inline_format(text)}</li>")

        # Empty line — paragraph break
        elif not stripped:
            flush_paragraph()

        # Regular text — accumulate paragraph
        else:
            in_paragraph = True
            paragraph_lines.append(inline_format(stripped))

    # Close any open elements
    flush_paragraph()
    if in_list:
        html_lines.append("</ul>")

    return "\n".join(html_lines)


# ============================================================================
# CHAPTER LOADING
# ============================================================================

def load_polished_chapters(workbook_dir: Path) -> list:
    """
    Load all approved chapters from w-polished/.
    Returns list of dicts sorted by chapter number.
    """
    polished_dir = workbook_dir / "w-polished"

    if not polished_dir.exists():
        error_exit(f"w-polished/ directory not found: {polished_dir}\nComplete human review first.")

    chapter_files = sorted(polished_dir.glob("chapter-*.md"))

    if not chapter_files:
        error_exit(f"No approved chapters found in {polished_dir}\nMove approved chapters from w-drafts/ to w-polished/")

    chapters = []
    for path in chapter_files:
        content = read_file(path)

        # Extract chapter number from filename
        num_match = re.search(r'chapter-(\d+)\.md', path.name)
        num = int(num_match.group(1)) if num_match else 0

        # Extract title from first H1
        title_match = re.search(r'^#\s+Chapter\s+\d+:\s*(.+)', content, re.M)
        if not title_match:
            title_match = re.search(r'^#\s+(.+)', content, re.M)
        title = title_match.group(1).strip() if title_match else f"Chapter {num}"

        chapters.append({
            "number": num,
            "title": title,
            "content": content,
            "path": path,
        })

    chapters.sort(key=lambda c: c["number"])
    log(f"Loaded {len(chapters)} approved chapters")
    return chapters


# ============================================================================
# HTML MANUSCRIPT ASSEMBLY
# ============================================================================

CSS = """
body {
    font-family: Georgia, 'Times New Roman', serif;
    font-size: 11pt;
    line-height: 1.6;
    max-width: 650px;
    margin: 0 auto;
    padding: 20px;
    color: #1a1a1a;
}
h1, h2, h3, h4 { font-family: 'Helvetica Neue', Arial, sans-serif; }
h1 { font-size: 2em; margin-top: 2em; }
h2 { font-size: 1.4em; margin-top: 1.8em; border-bottom: 1px solid #ddd; padding-bottom: 4px; }
h3 { font-size: 1.1em; margin-top: 1.5em; }
p { margin: 0.8em 0; text-align: justify; }
ul, ol { margin: 0.8em 0 0.8em 1.5em; }
li { margin: 0.3em 0; }
strong { font-weight: bold; }
em { font-style: italic; }
hr { border: none; border-top: 1px solid #ccc; margin: 2em 0; }
code { font-family: monospace; background: #f5f5f5; padding: 2px 4px; border-radius: 2px; }
.title-page { text-align: center; page-break-after: always; padding: 4em 0; }
.book-title { font-size: 2.5em; margin-bottom: 0.5em; }
.book-author { font-size: 1.3em; color: #555; }
.copyright-page { font-size: 0.85em; color: #555; margin: 2em 0; page-break-after: always; }
.toc { page-break-after: always; }
.toc ol { line-height: 2; }
.author-bio, .also-by { page-break-before: always; margin-top: 2em; }
.chapter-break { page-break-before: always; }
"""


def assemble_html(meta: dict, chapters: list, published_books: list) -> str:
    """Assemble full manuscript as HTML."""
    body_parts = []

    # Front matter
    body_parts.append(generate_title_page(meta))
    body_parts.append('<div class="chapter-break"></div>')
    body_parts.append(generate_copyright_page(meta))
    body_parts.append('<div class="chapter-break"></div>')
    body_parts.append(generate_toc_html(chapters))
    body_parts.append('<div class="chapter-break"></div>')

    # Chapters
    for i, ch in enumerate(chapters):
        chapter_id = f"chapter-{ch['number']:02d}"
        chapter_html = md_to_html(ch["content"], chapter_id)
        div_class = "chapter-break" if i > 0 else ""
        body_parts.append(f'<div class="{div_class}" id="{chapter_id}">')
        body_parts.append(chapter_html)
        body_parts.append("</div>")

    # Back matter
    body_parts.append('<div class="chapter-break"></div>')
    body_parts.append(generate_author_bio(meta["author"]))
    body_parts.append(generate_cta(meta, published_books))

    body = "\n".join(body_parts)

    return f"""<!DOCTYPE html>
<html lang="{meta['language']}">
<head>
<meta charset="UTF-8"/>
<meta name="author" content="{html.escape(meta['author'])}"/>
<meta name="description" content="{html.escape(meta['description'])}"/>
<title>{html.escape(meta['title'])}</title>
<style>
{CSS}
</style>
</head>
<body>
{body}
</body>
</html>"""


# ============================================================================
# EPUB GENERATION
# ============================================================================

def build_epub(meta: dict, chapters: list, output_dir: Path) -> Path:
    """Build EPUB using ebooklib (if available) or ebook-convert (Calibre)."""

    epub_path = output_dir / f"{meta['slug']}.epub"

    # Prefer Calibre ebook-convert if available (produces more reliable EPUB)
    ebook_convert = shutil.which("ebook-convert")
    if ebook_convert:
        log("Using Calibre ebook-convert for EPUB generation (preferred)")
        # Skip ebooklib entirely, go directly to calibre conversion
        pass
    else:
        # Try ebooklib first (only if Calibre not available)
        try:
            import ebooklib
            from ebooklib import epub

            book = epub.EpubBook()
            book.set_identifier(f"{meta['slug']}-{meta['date']}")
            book.set_title(meta["title"])
            book.set_language(meta["language"])
            book.add_author(meta["author"])

            # Add CSS
            css = epub.EpubItem(uid="style_nav", file_name="style/main.css",
                               media_type="text/css", content=CSS)
            book.add_item(css)

            # Add chapters
            epub_chapters = []
            for ch in chapters:
                chapter_id = f"chapter-{ch['number']:02d}"
                html_content = md_to_html(ch["content"], chapter_id)

                # Ensure content is non-empty for ebooklib
                if not html_content.strip():
                    html_content = f"<p>{ch['title']}</p>"

                c = epub.EpubHtml(
                    title=ch["title"],
                    file_name=f"{chapter_id}.xhtml",
                    lang=meta["language"],
                )
                c.content = f"""<?xml version='1.0' encoding='utf-8'?>
<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml">
<head><title>{ch['title']}</title>
<link rel="stylesheet" type="text/css" href="style/main.css"/>
</head>
<body>{html_content}</body></html>"""
                c.add_item(css)
                book.add_item(c)
                epub_chapters.append(c)

            book.toc = tuple(epub_chapters)
            book.add_item(epub.EpubNcx())
            book.add_item(epub.EpubNav())
            book.spine = ["nav"] + epub_chapters

            epub.write_epub(str(epub_path), book)
            log(f"EPUB built via ebooklib: {epub_path}")
            return epub_path

        except ImportError:
            log("ebooklib not found — cannot generate EPUB (no Calibre fallback)")
            return None
        except Exception as e:
            log(f"ebooklib failed ({e}) — cannot generate EPUB")
            return None

    # Calibre conversion path (if ebook_convert exists)
    html_path = output_dir / f"{meta['slug']}.html"
    if not html_path.exists():
        error_exit("HTML manuscript must be generated before EPUB via Calibre.")

    # cover_args already added later

    # Add cover if available
    cover_path = output_dir / "cover.jpg"
    cover_args = []
    if cover_path.exists():
        cover_args = ["--cover", str(cover_path)]

    cmd = [
        ebook_convert, str(html_path), str(epub_path),
        "--title", meta["title"],
        "--authors", meta["author"],
        "--language", meta["language"],
    ] + cover_args
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        log(f"WARNING: ebook-convert failed: {result.stderr[:200]}")
        return None

    log(f"EPUB built via Calibre: {epub_path}")
    return epub_path


# ============================================================================
# PDF GENERATION
# ============================================================================

def build_pdf(meta: dict, html_path: Path, output_dir: Path) -> Path:
    """Build PDF via WeasyPrint or Calibre fallback."""
    pdf_path = output_dir / f"{meta['slug']}.pdf"

    # Try WeasyPrint
    try:
        from weasyprint import HTML as WeasyHTML
        WeasyHTML(filename=str(html_path)).write_pdf(str(pdf_path))
        log(f"PDF built via WeasyPrint: {pdf_path}")
        return pdf_path
    except ImportError:
        log("weasyprint not found — trying ebook-convert")

    # Fallback: Calibre
    ebook_convert = shutil.which("ebook-convert")
    if not ebook_convert:
        log("WARNING: Neither weasyprint nor ebook-convert found. Skipping PDF.")
        log("  To fix: pip install weasyprint  OR  sudo apt install calibre")
        return None

    # Add cover if available
    cover_path = output_dir / "cover.jpg"
    cover_args = []
    if cover_path.exists():
        cover_args = ["--cover", str(cover_path)]

    cmd = [
        ebook_convert, str(html_path), str(pdf_path),
        "--title", meta["title"],
        "--authors", meta["author"],
    ] + cover_args
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        log(f"WARNING: ebook-convert PDF failed: {result.stderr[:200]}")
        return None

    log(f"PDF built via Calibre: {pdf_path}")
    return pdf_path


# ============================================================================
# KDP METADATA
# ============================================================================

def build_kdp_metadata(meta: dict, chapters: list, output_dir: Path) -> Path:
    """Write KDP-ready metadata JSON."""

    # BISAC category mapping — keyword-based heuristic
    title_lower = (meta["title"] + " " + meta.get("subtitle", "")).lower()
    bisac = "SEL027000"  # SELF-HELP / Personal Growth / General (default)

    if any(w in title_lower for w in ["sleep", "insomnia", "rest"]):
        bisac = "HEA006000"
    elif any(w in title_lower for w in ["gut", "nutrition", "diet", "eating", "food"]):
        bisac = "HEA017000"
    elif any(w in title_lower for w in ["weight", "walking", "fitness", "exercise", "workout"]):
        bisac = "HEA019000"
    elif any(w in title_lower for w in ["health", "fatigue", "chronic", "wellness", "medical"]):
        bisac = "HEA039000"
    elif any(w in title_lower for w in ["finance", "money", "budget", "invest", "income", "wealth"]):
        bisac = "BUS050000"
    elif any(w in title_lower for w in ["ai ", "artificial intelligence", "chatgpt", "machine learning"]):
        bisac = "COM014000"
    elif any(w in title_lower for w in ["security", "privacy", "network", "digital", "cyber", "hacking", "computer"]):
        bisac = "COM000000"
    elif any(w in title_lower for w in ["time management", "time blocking", "productivity", "adhd", "procrastin", "focus", "habit"]):
        bisac = "SEL016000"

    kdp_meta = {
        "title": meta["title"],
        "subtitle": meta.get("subtitle", ""),
        "author": meta["author"],
        "description": meta["description"],
        "keywords": meta["keywords"],
        "language": meta["language"],
        "publisher": meta["publisher"],
        "publication_date": meta["date"],
        "bisac_category": bisac,
        "adult_content": False,
        "enrollment": {
            "kdp_select": True,
            "print_replica": False,
        },
        "pricing": {
            "us_price": 9.99,
            "royalty_plan": "70%",
        },
        "chapter_count": len(chapters),
        "word_count_estimate": sum(len(ch["content"].split()) for ch in chapters),
        "generated": datetime.now().isoformat(),
    }

    path = output_dir / "kdp-metadata.json"
    path.write_text(json.dumps(kdp_meta, indent=2, ensure_ascii=False))
    log(f"KDP metadata written: {path}")
    return path


# ============================================================================
# VALIDATION REPORT
# ============================================================================

def generate_description_html(meta: dict) -> str:
    """
    Return KDP-ready HTML description.
    Priority: LLM-generated HTML (description_html field) > convert plain text.
    """
    # Use LLM-generated HTML if available
    if meta.get("description_html"):
        return meta["description_html"]

    # Fall back: convert plain text description to basic HTML
    raw = meta.get("description", f"A practical guide to {meta.get('title', 'this topic')}.")
    # Bold markdown **text** → <strong>text</strong>
    raw = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', raw)
    # Bullet lines starting with • or - → <ul><li>
    lines = raw.split("\n")
    result = []
    in_ul = False
    for line in lines:
        stripped = line.strip()
        if stripped.startswith(("• ", "- ", "* ")):
            if not in_ul:
                result.append("<ul>")
                in_ul = True
            result.append(f"<li>{stripped[2:].strip()}</li>")
        else:
            if in_ul:
                result.append("</ul>")
                in_ul = False
            if stripped:
                result.append(f"<p>{stripped}</p>")
    if in_ul:
        result.append("</ul>")
    return "\n".join(result)


# ── KDP Upload Kit ─────────────────────────────────────────────────────────────

# Suggested KDP categories for common BISAC codes
# These are VERIFIED paths that exist in KDP's category browser (as of 2025).
# Each niche has 3 options ordered by specificity: most-specific first.
# When uploading, navigate KDP's tree to confirm the path still exists.
BISAC_CATEGORIES = {
    # SELF-HELP / Productivity / Time Management
    "SEL027000": [
        "Kindle Store > Kindle eBooks > Self-Help > Personal Transformation",
        "Kindle Store > Kindle eBooks > Self-Help > Motivational",
        "Kindle Store > Kindle eBooks > Business & Money > Time Management",
    ],
    # SELF-HELP / Time Management
    "SEL016000": [
        "Kindle Store > Kindle eBooks > Self-Help > Time Management",
        "Kindle Store > Kindle eBooks > Business & Money > Time Management",
        "Kindle Store > Kindle eBooks > Self-Help > Personal Transformation",
    ],
    # HEALTH & FITNESS / Diseases / General  (chronic fatigue, medical conditions)
    "HEA039000": [
        "Kindle Store > Kindle eBooks > Health, Fitness & Dieting > Diseases & Physical Ailments > General",
        "Kindle Store > Kindle eBooks > Health, Fitness & Dieting > Mental Health",
        "Kindle Store > Kindle eBooks > Self-Help > Personal Transformation",
    ],
    # HEALTH & FITNESS / Exercise / General  (walking, fitness, weight loss)
    "HEA019000": [
        "Kindle Store > Kindle eBooks > Health, Fitness & Dieting > Exercise & Fitness",
        "Kindle Store > Kindle eBooks > Health, Fitness & Dieting > Weight Loss",
        "Kindle Store > Kindle eBooks > Health, Fitness & Dieting > Healthy Living",
    ],
    # HEALTH & FITNESS / Nutrition  (gut health, diet)
    "HEA017000": [
        "Kindle Store > Kindle eBooks > Health, Fitness & Dieting > Nutrition",
        "Kindle Store > Kindle eBooks > Health, Fitness & Dieting > Diets & Weight Loss > General",
        "Kindle Store > Kindle eBooks > Health, Fitness & Dieting > Healthy Living",
    ],
    # HEALTH & FITNESS / Sleep  (sleep books)
    "HEA006000": [
        "Kindle Store > Kindle eBooks > Health, Fitness & Dieting > Diseases & Physical Ailments > Sleep Disorders",
        "Kindle Store > Kindle eBooks > Health, Fitness & Dieting > Healthy Living",
        "Kindle Store > Kindle eBooks > Self-Help > Personal Transformation",
    ],
    # BUSINESS & ECONOMICS / Personal Finance
    "BUS050000": [
        "Kindle Store > Kindle eBooks > Business & Money > Personal Finance > General",
        "Kindle Store > Kindle eBooks > Business & Money > Entrepreneurship",
        "Kindle Store > Kindle eBooks > Business & Money > Skills > Decision-Making & Problem Solving",
    ],
    # COMPUTERS / Security
    "COM000000": [
        "Kindle Store > Kindle eBooks > Computers & Technology > Security & Encryption",
        "Kindle Store > Kindle eBooks > Computers & Technology > Networking & Cloud Computing",
        "Kindle Store > Kindle eBooks > Computers & Technology > Internet & Social Media > Online Safety & Privacy",
    ],
    # AI / Technology productivity
    "COM014000": [
        "Kindle Store > Kindle eBooks > Computers & Technology > Artificial Intelligence",
        "Kindle Store > Kindle eBooks > Computers & Technology > Software > Business",
        "Kindle Store > Kindle eBooks > Business & Money > Skills > Decision-Making & Problem Solving",
    ],
}

BISAC_LABELS = {
    "SEL027000": "SELF-HELP / Personal Growth / General",
    "SEL016000": "SELF-HELP / Time Management",
    "HEA039000": "HEALTH & FITNESS / Diseases / General",
    "HEA019000": "HEALTH & FITNESS / Exercise / General",
    "HEA017000": "HEALTH & FITNESS / Nutrition",
    "HEA006000": "HEALTH & FITNESS / Sleep",
    "BUS050000": "BUSINESS & ECONOMICS / Personal Finance / General",
    "COM000000": "COMPUTERS / Security / General",
    "COM014000": "COMPUTERS / Artificial Intelligence",
}


def write_upload_kit(meta: dict, outputs: dict, chapters: list, output_dir: Path) -> None:
    """
    Write kdp-upload-kit.txt — plain text, matches the format of published-book upload kits.
    Structure: KDP section → Gumroad section → Pre-Upload Checklist
    """
    total_words  = sum(len(ch["content"].split()) for ch in chapters)
    bisac        = meta.get("bisac_category", "SEL027000")
    categories   = BISAC_CATEGORIES.get(bisac, BISAC_CATEGORIES["SEL027000"])
    bisac_label  = BISAC_LABELS.get(bisac, bisac)
    description_html = generate_description_html(meta)
    description_plain = meta.get("description", "")
    keywords     = meta.get("keywords", [])
    while len(keywords) < 7:
        keywords.append("")

    gumroad_desc    = meta.get("gumroad_description", "")
    receipt_msg     = meta.get("receipt_message", "")
    docx_path       = str(outputs.get("docx", ""))
    cover_path      = str(output_dir / "cover.jpg") if (output_dir / "cover.jpg").exists() \
                      else "cover.jpg (run cover generator first)"
    title    = meta.get("title", "")
    subtitle = meta.get("subtitle", "")
    author   = meta.get("author", "William Archer")
    price    = meta.get("pricing", {}).get("us_price", 9.99)

    sep  = "=" * 60
    thin = "-" * 60

    lines = [
        sep,
        f"Upload Kit — {title}",
        sep,
        "",
        f"Title:      {title}",
        f"Subtitle:   {subtitle}",
        f"Author:     {author}",
        f"Word Count: {total_words:,}",
        f"Chapters:   {len(chapters)}",
        f"Book Files:",
        f"  Manuscript: {docx_path}",
        f"  Cover:      {cover_path}",
        "",
        sep,
        "KDP METADATA",
        sep,
        "",
        thin,
        "DESCRIPTION (plain text — paste as-is or use for reference)",
        thin,
        "",
        description_plain,
        "",
        thin,
        "DESCRIPTION (HTML — click </> in KDP description box, paste this)",
        thin,
        "",
        description_html,
        "",
        thin,
        "KEYWORDS (7 fields — one phrase each)",
        thin,
        "",
    ]

    for i, kw in enumerate(keywords[:7], 1):
        lines.append(f"  {i}. {kw}")

    lines += [
        "",
        thin,
        "CATEGORIES (navigate KDP's tree browser)",
        f"BISAC reference: {bisac_label}",
        thin,
        "",
    ]
    for i, cat in enumerate(categories[:3], 1):
        lines.append(f"  {i}. {cat}")

    lines += [
        "",
        thin,
        "PRICING",
        thin,
        "",
        f"  KDP eBook:   ${price:.2f}",
        f"  KDP Print:   $9.99",
        f"  Gumroad:     $9.99",
        "",
    ]

    if gumroad_desc:
        lines += [
            sep,
            "GUMROAD LISTING",
            sep,
            "",
            thin,
            "Gumroad Description",
            thin,
            "",
            gumroad_desc,
            "",
        ]

    if receipt_msg:
        lines += [
            thin,
            "Gumroad Custom Receipt Message",
            thin,
            "",
            receipt_msg,
            "",
        ]

    lines += [
        sep,
        "DRAFT2DIGITAL",
        sep,
        "",
        "Upload at: https://www.draft2digital.com/books",
        "Use EPUB file (D2D formats it — better than DOCX for their system)",
        f"  EPUB: {str(output_dir / 'book-*.epub') if not list(output_dir.glob('*.epub')) else str(next(output_dir.glob('*.epub'), ''))}",
        "",
        thin,
        "METADATA (same as KDP except where noted)",
        thin,
        "",
        f"  Title:           {title}",
        f"  Subtitle:        {subtitle}",
        f"  Author:          {author}",
        "  Language:        English",
        "  Series:          (leave blank)",
        "  Target Audience: Adult (18+)",
        "",
        "  Search Terms (type each phrase, press Enter after each):",
    ]

    for kw in keywords[:7]:
        if kw:
            lines.append(f"    {kw}")

    lines += [
        "",
        f"  BISAC: {bisac_label}",
        "",
        thin,
        "DESCRIPTION (plain text — D2D formats it, no HTML needed)",
        thin,
        "",
        description_plain,
        "",
        thin,
        "SHORT DESCRIPTION (1-2 sentences — shown on Smashwords store)",
        thin,
        "",
    ]

    # Generate a short description from the first sentence of description_plain
    if description_plain:
        first_sentence = description_plain.split(".")[0].strip("*# ") + "."
        lines.append(f"  {first_sentence}")
    else:
        lines.append(f"  A practical, evidence-based guide to {title.split(':')[0]}.")

    lines += [
        "",
        thin,
        "PRICING",
        thin,
        "",
        f"  USD Price: ${price:.2f}",
        "  Territorial prices: set CAD/GBP/AUD/EUR manually if desired",
        "",
        thin,
        "VENDORS (select all — Draft2Digital distributes to)",
        thin,
        "",
        "  Apple Books, Barnes & Noble, Kobo, Scribd, OverDrive,",
        "  Baker & Taylor, Smashwords, Tolino, and others",
        "  (use Toggle Button to select all at once)",
        "",
    ]

    lines += [
        sep,
        "PRE-UPLOAD CHECKLIST",
        sep,
        "",
        "KDP",
        f"  [ ] Manuscript uploaded (DOCX): {docx_path}",
        f"  [ ] Cover uploaded (JPG 1600x2560): {cover_path}",
        "  [ ] Description pasted (HTML source mode)",
        "  [ ] 7 keywords entered",
        "  [ ] 3 categories selected",
        f"  [ ] Price set (${price:.2f} eBook)",
        "  [ ] Preview reviewed in KDP previewer",
        "  [ ] Adult content: No",
        "  [ ] Publishing rights confirmed",
        "",
        "Gumroad",
        "  [ ] Product page created",
        "  [ ] PDF + EPUB uploaded",
        "  [ ] Description set",
        "  [ ] Cover/thumbnail set",
        "  [ ] Price set ($9.99)",
        "  [ ] Button message set",
        "  [ ] Custom receipt message set",
        "",
        "Post-Upload",
        "  [ ] KDP preview approved",
        "  [ ] Set 7-day performance check reminder",
        "  [ ] Set 30-day performance check reminder",
        "",
        "Draft2Digital",
        "  [ ] EPUB uploaded",
        "  [ ] Metadata entered (title, subtitle, author, search terms, BISAC)",
        "  [ ] Description entered (plain text)",
        "  [ ] Layout reviewed (no duplicate front matter)",
        "  [ ] Price set",
        "  [ ] All vendors selected",
        "  [ ] Submitted",
        "",
        sep,
        "Remember: 80% shipped beats 100% perfect.",
        "Get this live, then iterate based on feedback.",
        sep,
    ]

    kit_path = output_dir / "kdp-upload-kit.txt"
    kit_path.write_text("\n".join(lines), encoding="utf-8")
    log(f"KDP upload kit written: {kit_path}")


def write_package_report(meta: dict, outputs: dict, chapters: list, output_dir: Path) -> None:
    """Write a package validation report."""
    total_words = sum(len(ch["content"].split()) for ch in chapters)
    lines = [
        f"# Package Report: {meta['title']}",
        f"<!-- AUTO-GENERATED: {datetime.now().isoformat(timespec='seconds')} -->",
        "",
        f"- Author: {meta['author']}",
        f"- Chapters: {len(chapters)}",
        f"- Total words: {total_words:,}",
        f"- Date: {meta['date']}",
        "",
        "## Output Files",
    ]

    for fmt, path in outputs.items():
        status = "✓" if path and Path(path).exists() else "✗ (not generated)"
        lines.append(f"- {fmt.upper()}: {path} {status}")

    lines += [
        "",
        "## Chapters Included",
    ]
    for ch in chapters:
        words = len(ch["content"].split())
        lines.append(f"- Chapter {ch['number']:02d}: {ch['title']} ({words:,} words)")

    lines += [
        "",
        "## Validation",
        f"- Chapter count: {len(chapters)} {'✓' if len(chapters) >= 10 else '✗ (expected 10-12)'}",
        f"- Total words: {total_words:,} {'✓' if total_words >= 20000 else '✗ (low)'}",
        "",
        "## Next Steps",
        "1. Generate cover: python3 cover_generator.py --book-dir <workbook>",
        "2. Open kdp-upload-kit.txt — follow the steps in order",
        "3. Upload DOCX + cover at kdp.amazon.com",
    ]

    report_path = output_dir / "package-report.md"
    report_path.write_text("\n".join(lines))
    log(f"Package report: {report_path}")


# ============================================================================
# MAIN
# ============================================================================

def main():
    parser = argparse.ArgumentParser(
        description="Packager Agent — Ebook Factory Phase 3",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python3 packager.py
  python3 packager.py --book-dir ~/.hermes/ebook-factory/workbooks/book-my-topic/
  python3 packager.py --author "William Archer" --title "My Book Title"
  python3 packager.py --formats epub,html
        """
    )
    parser.add_argument("--book-dir", type=str, default=None,
                        help="Workbook directory (auto-detects latest if not set)")
    parser.add_argument("--author", type=str, default=None,
                        help=f"Author name (default: {DEFAULT_AUTHOR})")
    parser.add_argument("--title", type=str, default=None,
                        help="Override book title")
    parser.add_argument("--formats", type=str, default="html,epub,pdf",
                        help="Comma-separated output formats: html,epub,pdf (default: all)")

    args = parser.parse_args()

    # Determine workbook
    if args.book_dir:
        workbook_dir = Path(args.book_dir)
        if not workbook_dir.exists():
            error_exit(f"Book directory not found: {workbook_dir}")
    else:
        workbook_dir = find_latest_workbook()
        log(f"Auto-detected workbook: {workbook_dir.name}")

    log(f"Packaging: {workbook_dir}")

    # Parse formats
    formats = [f.strip().lower() for f in args.formats.split(",")]

    # Paths
    outline_path = workbook_dir / "01_outline.md"
    output_dir = workbook_dir / "output"
    output_dir.mkdir(parents=True, exist_ok=True)

    # Get list of published books for CTA
    published_books = []
    pub_dir = Path.home() / "books" / "factory" / "references" / "published-books"
    if pub_dir.exists():
        for d in pub_dir.iterdir():
            if d.is_dir():
                # Extract book name from dir name "1. Book Title" -> "Book Title"
                name = re.sub(r'^\d+\.\s*', '', d.name)
                published_books.append(name)

    # Load metadata
    log_section("Step 1: Loading metadata")
    meta = extract_book_metadata(workbook_dir, outline_path, args)
    log(f"Title: {meta['title']}")
    log(f"Author: {meta['author']}")
    log(f"Keywords: {meta['keywords']}")

    # Load chapters
    log_section("Step 2: Loading approved chapters")
    chapters = load_polished_chapters(workbook_dir)
    for ch in chapters:
        words = len(ch["content"].split())
        log(f"  Chapter {ch['number']:02d}: {ch['title']} ({words:,} words)")

    outputs = {}

    # HTML
    if "html" in formats:
        log_section("Step 3: Assembling HTML manuscript")
        html = assemble_html(meta, chapters, published_books)
        html_path = output_dir / f"{meta['slug']}.html"
        write_file_atomic(html_path, html)
        outputs["html"] = html_path

    # EPUB
    if "epub" in formats:
        log_section("Step 4: Building EPUB")
        epub_path = build_epub(meta, chapters, output_dir)
        outputs["epub"] = epub_path

    # PDF
    if "pdf" in formats:
        log_section("Step 5: Building PDF")
        html_path = output_dir / f"{meta['slug']}.html"
        if html_path.exists():
            pdf_path = build_pdf(meta, html_path, output_dir)
            outputs["pdf"] = pdf_path
        else:
            log("WARNING: HTML not generated — skipping PDF")

    # KDP metadata
    log_section("Step 6: Writing KDP metadata")
    kdp_path = build_kdp_metadata(meta, chapters, output_dir)
    outputs["kdp-metadata"] = kdp_path

    # DOCX — KDP's most reliable upload format (pandoc)
    log_section("Step 6b: Building DOCX (KDP upload format)")
    html_path = output_dir / f"{meta['slug']}.html"
    docx_path = output_dir / f"{meta['slug']}.docx"
    pandoc = shutil.which("pandoc")
    if pandoc and html_path.exists():
        try:
            # NOTE: Do NOT pass --metadata title/author or --toc here.
            # --metadata title/author makes pandoc auto-generate a Title+Author
            # block that duplicates the HTML title-page div.
            # --toc creates a Word TOC field code that (a) appears at the very
            # start of the document (before the title page), and (b) shows as an
            # empty placeholder in LibreOffice/ONLYOFFICE that requires manual
            # "Update Field" to populate. The HTML already has a properly
            # formatted TOC with chapter links that pandoc will convert normally.
            result = subprocess.run(
                [
                    pandoc,
                    str(html_path),
                    "-o", str(docx_path),
                ],
                capture_output=True, text=True, timeout=60
            )
            if result.returncode == 0 and docx_path.exists():
                # Post-process: fix title page formatting in the DOCX
                try:
                    _fix_docx_title_page(docx_path, meta)
                except Exception as e:
                    log(f"WARNING: DOCX title page fix failed: {e}")
                log(f"DOCX built via pandoc: {docx_path}")
                outputs["docx"] = docx_path
            else:
                log(f"WARNING: pandoc DOCX failed: {result.stderr[:200]}")
        except Exception as e:
            log(f"WARNING: DOCX generation failed: {e}")
    else:
        if not pandoc:
            log("WARNING: pandoc not found — skipping DOCX. Install with: sudo apt install pandoc")
        else:
            log("WARNING: HTML not found — skipping DOCX")

    # Package report
    log_section("Step 7: Writing package report")
    write_package_report(meta, outputs, chapters, output_dir)

    # KDP upload kit
    log_section("Step 7b: Writing KDP upload kit")
    write_upload_kit(meta, outputs, chapters, output_dir)

    # Summary
    log_section("PACKAGING COMPLETE")
    total_words = sum(len(ch["content"].split()) for ch in chapters)
    log(f"Book: {meta['title']}")
    log(f"Chapters: {len(chapters)}")
    log(f"Total words: {total_words:,}")
    log(f"Output: {output_dir}")
    log("")
    log("Files:")
    for fmt, path in outputs.items():
        if path and Path(str(path)).exists():
            log(f"  ✓ {fmt}: {path}")
        elif path:
            log(f"  ✗ {fmt}: FAILED")
    log("")
    log("Next: Upload to Amazon KDP.")
    log("  → Open output/kdp-upload-kit.txt for a complete copy-paste guide.")
    log("  → KDP UPLOAD ORDER (most reliable first):")
    log("  1. DOCX  — most reliable, KDP converts cleanly")
    log("  2. EPUB  — run through eBook-Standardization-Toolkit first")
    log("  3. HTML  — fallback if others fail")
    sys.exit(0)


if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\nInterrupted.", flush=True)
        sys.exit(130)
