#!/usr/bin/env python3
"""Convert EPUB to DOCX and apply formatting fixes for the 7 published books
that only have EPUB (no DOCX).

Steps:
1. pandoc epub -> docx
2. Apply title page + TOC + chapter post-processing

Usage:
  python3 epub_to_docx.py
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from pathlib import Path
import re
import shutil
import subprocess

WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def fix_docx(docx_path):
    """Apply full title page + TOC + chapter formatting."""
    doc = Document(str(docx_path))
    paragraphs = doc.paragraphs

    # Phase 1: Style the title page
    title_done = False
    subtitle_done = False
    author_done = False
    
    for i, p in enumerate(paragraphs[:6]):
        text = p.text.strip()
        
        if p.style.name == 'Title' and not title_done:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(28)
                run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
            p.paragraph_format.space_before = Pt(72)
            p.paragraph_format.space_after = Pt(12)
            title_done = True
            print('  Styled title')

        elif p.style.name == 'Subtitle' and not subtitle_done:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(16)
                run.font.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            p.paragraph_format.space_after = Pt(24)
            subtitle_done = True
            print('  Styled subtitle')

        elif p.style.name == 'Author' and not author_done:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            p.paragraph_format.space_before = Pt(36)
            p.paragraph_format.space_after = Pt(6)
            # Page break after author
            run_elem = p._element.makeelement(qn('w:br'), {qn('w:type'): 'page'})
            new_run = p.add_run()
            new_run._element.append(run_elem)
            author_done = True
            print('  Styled author + page break')

        elif p.style.name == 'Date':
            # Remove Date line (redundant with copyright)
            parent = p._element.getparent()
            parent.remove(p._element)
            print('  Removed Date line')

    # Phase 2: Find and style TOC + add chapter page breaks + bookmarks
    in_toc = False
    toc_count = 0
    paragraphs = doc.paragraphs

    for i, p in enumerate(paragraphs):
        text = p.text.strip()

        # Detect TOC heading
        if 'Table of Contents' in text and 'Heading' in p.style.name:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(20)
                run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(18)
            in_toc = True
            toc_count = 0
            print('  Styled TOC heading')
            continue

        # Style TOC entries
        if in_toc and p.style.name in ('Compact', 'TOCEntry'):
            toc_count += 1
            for r_elem in p._element.iter('{' + WNS + '}r'):
                rpr = r_elem.find('{' + WNS + '}rPr')
                if rpr is None:
                    rpr = r_elem.makeelement('{' + WNS + '}rPr', {})
                    r_elem.insert(0, rpr)
                sz = rpr.find('{' + WNS + '}sz')
                if sz is None:
                    sz = r_elem.makeelement('{' + WNS + '}sz', {})
                    rpr.append(sz)
                sz.set('{' + WNS + '}val', '26')
                szCs = rpr.find('{' + WNS + '}szCs')
                if szCs is None:
                    szCs = r_elem.makeelement('{' + WNS + '}szCs', {})
                    rpr.append(szCs)
                szCs.set('{' + WNS + '}val', '26')
                color = rpr.find('{' + WNS + '}color')
                if color is None:
                    color = r_elem.makeelement('{' + WNS + '}color', {})
                    rpr.append(color)
                color.set('{' + WNS + '}val', '333333')
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.space_before = Pt(2)
            continue

        if in_toc and p.style.name not in ('Compact', 'TOCEntry', 'Body Text'):
            in_toc = False
            # Page break after last TOC entry
            if i > 0 and paragraphs[i-1].style.name in ('Compact', 'TOCEntry'):
                prev = paragraphs[i-1]
                run_elem = prev._element.makeelement(qn('w:br'), {qn('w:type'): 'page'})
                new_run = prev.add_run()
                new_run._element.append(run_elem)
                print('  Page break after TOC (%d entries)' % toc_count)

        # Chapter headings: page break + bookmark
        if 'Chapter' in text and 'Heading' in p.style.name:
            p.paragraph_format.page_break_before = True
            ns = {'w': WNS}
            existing_bms = p._element.findall('.//w:bookmarkStart', ns)
            if not existing_bms:
                ch_match = re.search(r'Chapter\s+(\d+)', text)
                if ch_match:
                    ch_num = int(ch_match.group(1))
                    bm_name = 'chapter-%02d' % ch_num
                    bm_id = 100 + ch_num
                    bm_start = p._element.makeelement('{' + WNS + '}bookmarkStart', {
                        '{' + WNS + '}id': str(bm_id),
                        '{' + WNS + '}name': bm_name,
                    })
                    p._element.insert(0, bm_start)
                    bm_end = p._element.makeelement('{' + WNS + '}bookmarkEnd', {
                        '{' + WNS + '}id': str(bm_id),
                    })
                    p._element.append(bm_end)
                    print('  Page break + bookmark %s' % bm_name)

    doc.save(str(docx_path))
    print('  Saved: ' + str(docx_path))


# Books that need EPUB -> DOCX conversion
BOOKS = [
    '2. The 80-20 Guide to Getting More Done',
    '3. Weekly Review Systems for Busy People',
    '4. Building a Second Brain on a Budget',
    '6. Walking for Weight Loss',
    '7. AI Tools for Everyday Productivity',
    '9. Home Network Security',
    '12. Low-Income Chronic Fatigue Management',
]

PUB_DIR = Path('/home/bookforge/books/factory/references/published-books')

for book_dir_name in BOOKS:
    book_dir = PUB_DIR / book_dir_name
    if not book_dir.exists():
        print('SKIP: %s (not found)' % book_dir_name)
        continue

    # Find EPUB
    epubs = list(book_dir.glob('*.epub'))
    if not epubs:
        print('SKIP: %s (no EPUB)' % book_dir_name)
        continue

    epub_path = epubs[0]
    # Derive DOCX name from book dir
    docx_name = book_dir_name.split('. ', 1)[1].lower().replace(' ', '-').replace(',', '') + '.docx'
    docx_path = book_dir / docx_name

    print('=== %s ===' % book_dir_name)
    print('  EPUB: %s' % epub_path.name)
    print('  DOCX: %s' % docx_name)

    # Step 1: pandoc EPUB -> DOCX
    result = subprocess.run(
        ['pandoc', str(epub_path), '-o', str(docx_path)],
        capture_output=True, text=True, timeout=120
    )
    if result.returncode != 0:
        print('  PANDOC FAILED: %s' % result.stderr[:200])
        continue
    print('  Pandoc conversion OK')

    # Step 2: Apply formatting fixes
    try:
        fix_docx(docx_path)
    except Exception as e:
        print('  FIX FAILED: %s' % str(e))
    print()
